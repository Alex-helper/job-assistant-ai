"""
求职助手 AI - FastAPI 后端
本地：uvicorn server:app --host 127.0.0.1 --port 8765
公网：由平台注入 PORT；设置 PUBLIC_MODE=1
"""

import os
import json
import uuid
import re
import secrets
from pathlib import Path
from typing import List, Optional
from datetime import datetime, timedelta

import base64
import io
import subprocess
import tempfile
import sys
from fastapi import FastAPI, File, UploadFile, Request, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

BASE_DIR = Path(__file__).resolve().parent
ENV_FILE = BASE_DIR / ".env"
# utf-8-sig：兼容编辑器写入的 BOM，避免键名变成 \ufeffOPENAI_API_KEY
load_dotenv(dotenv_path=ENV_FILE, override=True, encoding="utf-8-sig")

from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage

from rag_store import get_rag
from job_boards import (
    extract_keywords_heuristic,
    hits_to_dict,
    hits_to_markdown,
    list_company_presets,
    list_fame_tiers,
    list_sources,
    search_jobs,
)

# ── 公网安全开关 ───────────────────────────────────────────
def _env_truthy(name: str, default: str = "0") -> bool:
    return os.getenv(name, default).strip().lower() in ("1", "true", "yes", "on")

# 默认本地开发关闭公网限制；Render 等平台请设 PUBLIC_MODE=1
PUBLIC_MODE = _env_truthy("PUBLIC_MODE", "0")
ENABLE_RUN_CODE = _env_truthy("ENABLE_RUN_CODE", "0" if PUBLIC_MODE else "1")
ADMIN_TOKEN = os.getenv("ADMIN_TOKEN", "").strip()

STATIC_DIR = BASE_DIR / "static"

def _cors_origins() -> list:
    raw = os.getenv("CORS_ORIGINS", "").strip()
    if raw:
        return [o.strip() for o in raw.split(",") if o.strip()]
    if PUBLIC_MODE:
        # 同源部署时浏览器不跨域；收紧为拒绝任意 Origin 预检滥用
        return []
    return ["*"]

# ── 应用初始化 ─────────────────────────────────────────────
app = FastAPI(title="求职助手 API", docs_url=None, redoc_url=None)

_origins = _cors_origins()
app.add_middleware(
    CORSMiddleware,
    allow_origins=_origins if _origins else [],
    allow_credentials=False,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)

def _require_admin_token(request: Request, password: Optional[str] = None) -> None:
    """公网模式：必须带 ADMIN_TOKEN（Header / query）；本地仍可用支付配置里的 password。"""
    if not PUBLIC_MODE:
        return
    if not ADMIN_TOKEN:
        raise HTTPException(status_code=503, detail="未配置 ADMIN_TOKEN，管理接口已禁用")
    auth = request.headers.get("Authorization", "")
    bearer = auth[7:].strip() if auth.lower().startswith("bearer ") else ""
    token = (
        bearer
        or request.headers.get("X-Admin-Token", "").strip()
        or request.query_params.get("admin_token", "").strip()
        or (password or "").strip()
    )
    if not token or not secrets.compare_digest(token, ADMIN_TOKEN):
        raise HTTPException(status_code=401, detail="未授权")

# ── 工具函数 ───────────────────────────────────────────────
def get_llm(temperature: float = 0.3, max_tokens: Optional[int] = None):
    kwargs = {
        "model": os.getenv("MODEL_NAME", "gpt-4o-mini"),
        "temperature": temperature,
        "streaming": True,
    }
    if max_tokens:
        kwargs["max_tokens"] = max_tokens
    return ChatOpenAI(**kwargs)

async def sse_stream(
    messages_or_prompt,
    temperature: float = 0.3,
    max_tokens: Optional[int] = None,
    meta: Optional[dict] = None,
):
    """将 LLM 响应转为 SSE 流；可选先推送 meta（如 RAG sources）。"""
    if meta:
        yield f"data: {json.dumps({'meta': meta}, ensure_ascii=False)}\n\n"
    llm = get_llm(temperature=temperature, max_tokens=max_tokens)
    try:
        async for chunk in llm.astream(messages_or_prompt):
            if chunk.content:
                data = json.dumps({"content": chunk.content}, ensure_ascii=False)
                yield f"data: {data}\n\n"
    except Exception as e:
        err = json.dumps({"error": str(e)}, ensure_ascii=False)
        yield f"data: {err}\n\n"
    yield "data: [DONE]\n\n"

def stream_resp(
    prompt_or_msgs,
    temperature: float = 0.3,
    max_tokens: Optional[int] = None,
    meta: Optional[dict] = None,
):
    return StreamingResponse(
        sse_stream(prompt_or_msgs, temperature=temperature, max_tokens=max_tokens, meta=meta),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )

# ── 请求模型 ───────────────────────────────────────────────
class JDReq(BaseModel):
    jd: str

class MatchReq(BaseModel):
    jd: str
    resume: str
    use_rag: bool = True


class CoverLetterReq(BaseModel):
    jd: str
    resume: str
    name: str = "求职者"
    company: str = "贵公司"

class InterviewReq(BaseModel):
    jd: str
    resume: str = ""


class JobMatchReq(BaseModel):
    resume: str
    city: str = ""
    role_hint: str = ""
    # large=大厂 mid=中厂 small=小厂（按品牌知名度）
    fame_tiers: List[str] = ["large", "mid", "small"]
    keywords: str = ""  # 可选，逗号分隔；空则从简历抽取
    companies: List[str] = []  # 可选目标公司（勾选/自填）
    max_hits: int = 20


class ProjectRecReq(BaseModel):
    jd: str
    background: str = ""
    preference: str = ""  # 技术偏好 / 工期约束（可选）
    python_version: str = ""
    cuda_version: str = ""
    gpu_model: str = ""
    include_finetune: bool = True
    include_pretrain: bool = False
    include_interview_qa: bool = False
    include_multimodal: bool = False
    include_github: bool = True
    include_demo: bool = True
    # 兼容旧前端：若仍传 include_deploy，则同时开/关 GitHub+Demo
    include_deploy: Optional[bool] = None
    # 推理 API：cn=国内 / overseas=国外 / local=纯本地
    llm_region: str = "cn"
    llm_provider: str = ""
    llm_model: str = ""
    # detailed=详细版（默认） / brief=简短版
    detail_level: str = "detailed"
    # 生成几个项目：1/2/3
    project_count: int = 3
    # 项目质量定位：主推 / 备选 / 加分（数量应与 project_count 一致）
    project_tiers: List[str] = ["主推", "备选", "加分"]
    # 项目完成时间（天）：1 / 3 / 7 / 15 / 30
    timeline_days: int = 7
    # 项目背景：school=学校级 / enterprise=企业级
    project_scale: str = "school"
    # 项目形式：web=网站 / extension=浏览器插件 / desktop=电脑端App / mobile=移动端App
    project_form: str = "web"
    # 是否差异化选题：True=避免与市面 Demo 同质化 / False=允许客服/问答等常规项目
    differentiate: bool = True


class ProjectDraftReq(BaseModel):
    content: str
    jd: str = ""
    message: str = ""
    include_finetune: bool = True
    include_pretrain: bool = False
    include_interview_qa: bool = False
    include_multimodal: bool = False
    include_github: bool = True
    include_demo: bool = True
    include_deploy: Optional[bool] = None
    detail_level: str = "detailed"
    project_count: int = 3
    project_tiers: List[str] = ["主推", "备选", "加分"]
    timeline_days: int = 7
    project_scale: str = "school"
    project_form: str = "web"
    differentiate: bool = True
    # 项目推荐 Agent 是否启用飞书知识库 RAG（不影响 /api/project-recommend 原提示词）
    use_rag: bool = True


class ExportDocReq(BaseModel):
    content: str
    format: str = "docx"  # docx | pdf
    title: str = "项目推荐方案"


class ResumeTemplateReq(BaseModel):
    jd: str = ""
    resume: str = ""
    style: str = "star"


class ResumeFillReq(BaseModel):
    jd: str
    resume: str
    template: str
    use_rag: bool = True


class StarOptimizeReq(BaseModel):
    """简历优化栏 · STAR 项目专项改写。"""
    project: str  # 原项目经历 / STAR 草稿
    jd: str = ""
    resume: str = ""  # 可选上下文
    use_rag: bool = True
    enhance_stack: bool = True  # 允许适当增补技术栈
    domain_bg: str = ""  # 目标业务背景，如「法律」「医疗」「电商」
    domain_verdict: str = ""  # 检测结果：ok / risky / reject
    force_domain: bool = False  # 检测为不建议时仍强制替换


class StarDomainCheckReq(BaseModel):
    """检测项目背景是否适合替换到指定领域。"""
    project: str
    domain_bg: str
    use_rag: bool = True


class ChatMsg(BaseModel):
    role: str
    content: str


class ChatReq(BaseModel):
    message: str
    history: List[ChatMsg] = []

class ConfigReq(BaseModel):
    api_key: str
    base_url: str = "https://api.deepseek.com/v1"
    model: str = "deepseek-chat"
    provider_id: str = "deepseek"
    region: str = "cn"


PROVIDERS_CATALOG = Path.home() / ".cursor" / "skills" / "global-api-env" / "providers.json"


@app.get("/api/providers-catalog")
async def providers_catalog():
    if not PROVIDERS_CATALOG.exists():
        return JSONResponse({"error": "providers catalog missing"}, status_code=404)
    return json.loads(PROVIDERS_CATALOG.read_text(encoding="utf-8"))

# ── 页面路由 ───────────────────────────────────────────────
@app.get("/")
async def root():
    html_file = STATIC_DIR / "index.html"
    return HTMLResponse(
        content=html_file.read_text(encoding="utf-8"),
        headers={
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )

@app.get("/static/{filename}")
async def static_file(filename: str):
    from fastapi.responses import FileResponse
    f = STATIC_DIR / filename
    if not f.exists():
        return JSONResponse({"error": "not found"}, status_code=404)
    return FileResponse(str(f))

# ── 系统接口 ───────────────────────────────────────────────
@app.get("/api/health")
async def health():
    key = os.getenv("OPENAI_API_KEY", "")
    configured = bool(key) and not key.startswith("sk-xxx")
    masked = ("*" * (len(key) - 4) + key[-4:]) if configured else ""
    return {
        "configured": configured,
        "key_hint": masked,
        "model": os.getenv("MODEL_NAME", "gpt-4o-mini"),
        "base_url": os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1"),
        "public_mode": PUBLIC_MODE,
        "run_code_enabled": ENABLE_RUN_CODE,
    }

@app.get("/api/config")
async def get_config():
    """公网不返回明文 Key；本地开发仍可回填到设置面板。"""
    key = os.getenv("OPENAI_API_KEY", "")
    configured = bool(key) and not key.startswith("sk-xxx")
    provider_id = os.getenv("API_PROVIDER_ID", "deepseek")
    region = os.getenv("API_REGION", "cn")
    base_default = "https://api.deepseek.com/v1"
    model_default = "deepseek-chat"
    if PUBLIC_MODE:
        masked = ("*" * max(0, len(key) - 4) + key[-4:]) if configured else ""
        return {
            "api_key": "",
            "key_hint": masked,
            "configured": configured,
            "base_url": os.getenv("OPENAI_BASE_URL", base_default),
            "model": os.getenv("MODEL_NAME", model_default),
            "provider_id": provider_id,
            "region": region,
            "public_mode": True,
            "writable": False,
        }
    return {
        "api_key": key,
        "base_url": os.getenv("OPENAI_BASE_URL", base_default),
        "model": os.getenv("MODEL_NAME", model_default),
        "provider_id": provider_id,
        "region": region,
        "public_mode": False,
        "writable": True,
    }

def _upsert_env(path: Path, updates: dict) -> None:
    """原地更新 .env 键值，避免 set_key 追加重复行/多余引号。"""
    text = path.read_text(encoding="utf-8-sig") if path.exists() else ""
    lines = text.splitlines()
    seen = set()
    out = []
    for line in lines:
        raw = line.strip()
        if not raw or raw.startswith("#") or "=" not in raw:
            out.append(line)
            continue
        key = raw.split("=", 1)[0].strip().lstrip("\ufeff")
        if key in updates:
            if key not in seen:
                out.append(f"{key}={updates[key]}")
                seen.add(key)
            continue
        out.append(line)
    for key, val in updates.items():
        if key not in seen:
            out.append(f"{key}={val}")
    # 无 BOM 写入，避免 OPENAI_API_KEY 被读成 \ufeffOPENAI_API_KEY
    path.write_bytes(("\n".join(out).rstrip() + "\n").encode("utf-8"))


@app.post("/api/config")
async def save_config(req: ConfigReq):
    """将配置写入 .env 文件并立即生效。公网环境禁用（Key 仅通过平台环境变量注入）。"""
    if PUBLIC_MODE:
        raise HTTPException(
            status_code=403,
            detail="公网模式禁止通过接口修改 API Key，请用本地启动脚本，或直接编辑 .env",
        )
    _upsert_env(ENV_FILE, {
        "API_PROVIDER_ID": (req.provider_id or "deepseek").strip(),
        "API_REGION": (req.region or "cn").strip(),
        "OPENAI_API_KEY": req.api_key.strip(),
        "OPENAI_BASE_URL": (req.base_url or "https://api.deepseek.com/v1").strip(),
        "MODEL_NAME": (req.model or "deepseek-chat").strip(),
    })
    load_dotenv(dotenv_path=str(ENV_FILE), override=True, encoding="utf-8-sig")
    return {"ok": True}

# ── AI 功能接口 ────────────────────────────────────────────
@app.post("/api/analyze-jd")
async def analyze_jd(req: JDReq):
    prompt = f"""请对以下招聘 JD 进行结构化分析，使用 Markdown 格式输出中文结果。

## JD 原文
{req.jd}

---

## 📌 岗位定位
（一句话概括岗位核心价值）

## 🎯 核心职责（Top 5）
1. 
2. 
...

## ✅ 硬性要求
**学历/经验：**  
**必备技术栈：**  
**工具/平台：**  

## ⭐ 加分项
- 

## 💬 软技能要求
- 

## 🔑 面试关键词
> 建议在面试中主动提及以下词汇

`关键词1` `关键词2` `关键词3` ..."""
    return stream_resp(prompt)

@app.post("/api/evaluate-match")
async def evaluate_match(req: MatchReq):
    prompt = f"""你是资深 HR，请评估以下简历与岗位的匹配度，使用 Markdown 格式输出中文结果。

## 简历内容
{req.resume}

## 岗位 JD
{req.jd}

---

## 🎯 综合匹配分：__/100

## ✅ 优势项
（简历与 JD 高度匹配的亮点）

## ❌ 差距项
（JD 要求但简历缺失或薄弱的地方）

## 📋 提升优先级

| 优先级 | 改进项 | 具体建议 |
|:------:|--------|---------|
| 🔴 高  | ...    | ...     |
| 🟡 中  | ...    | ...     |
| 🟢 低  | ...    | ...     |

## 💡 总体建议
（2-3 句综合评价与行动建议）"""
    return stream_resp(prompt)

DEFAULT_RESUME_TEMPLATE = """# {{姓名}}｜{{求职意向}}

## 个人信息
- 手机：{{手机}}
- 邮箱：{{邮箱}}
- 城市：{{城市}}
- GitHub / 作品集：{{链接}}
- 实习意向：可实习 **6 个月及以上**；可 **一周内到岗**；出勤 **每周 5 天（全勤）**

## 教育经历
- {{学校}}｜{{专业}}｜{{学历}}｜{{起止时间}}
  - 主修课程 / 成绩：{{补充}}

## 专业技能
- 编程语言：{{语言}}
- 框架与工具：{{框架}}
- AI / 数据：{{AI技能}}
- 其他：{{其他技能}}

## 项目经历

### {{项目名称一}}（{{项目时间}}）
**技术栈：** {{技术栈列表，用 / 分隔}}

**项目介绍（STAR）：**
- **S（情境）：** {{业务背景与约束}}
- **T（任务）：** {{你的目标与职责}}
- **A（行动）：** {{关键做法与技术判断；正文中出现的技术名请用 **加粗**}}
- **R（结果）：** {{可量化结果；没有数字写【待补充】}}

### {{项目名称二}}（{{项目时间}}）
**技术栈：** {{技术栈列表，用 / 分隔}}

**项目介绍（STAR）：**
- **S（情境）：** ...
- **T（任务）：** ...
- **A（行动）：** ...
- **R（结果）：** ...

## 获奖经历
- {{奖项名称}}｜{{颁发机构}}｜{{获奖时间}}
  - 简要说明：{{含金量/排名/与岗位相关点}}
- {{奖项名称二}}｜{{颁发机构}}｜{{获奖时间}}
"""

RESUME_FORMAT_RULES = """
## 简历排版硬性规则（必须遵守）
1. 模块顺序固定：个人信息 → 教育经历 → 专业技能 → 项目经历 → 获奖经历（可按模板增减小节，但不得打乱上述主序）。
2. 每个项目必须先写一行：**技术栈：** xxx / yyy / zzz（放在项目介绍上面）。
3. 项目介绍必须用 STAR：S / T / A / R 四行（可用加粗标签）。
4. 项目介绍正文里出现的技术名、框架名、模型名必须用 Markdown **加粗**（例如 **PyTorch**、**FastAPI**、**LoRA**）。
5. 获奖经历放在项目经历之后；无奖项时写「暂无」或【待补充】，不要编造。
6. **个人信息必须包含实习/到岗信息（默认写上，除非原材料明确冲突）**：
   - 可实习 **6 个月及以上**
   - 可 **一周内到岗**
   - 出勤 **每周 5 天（全勤到岗）**
   写成一行「实习意向：…」放在个人信息内。
7. 禁止编造获奖/实习单位等事实；工程指标须贴合项目规模、可面试追问，并尽量写「基线方法 → 改进方法 → 相对变化」。
8. 输出完整可投递 Markdown 简历，不要只给修改建议（除非用户只要建议）。
"""

STAR_PROJECT_RULES = """
## STAR 项目改写规则（必须遵守）
1. **保留原项目主线**：业务场景、你的角色、核心成果方向不要换成另一个无关项目。
2. **可适当增强**：在原项目合理范围内增补技术栈（优先来自飞书知识库中较新、可落地的方法），写进 **技术栈：** 与 STAR-A；增补要能讲清「为什么用、解决什么问题」，禁止堆砌无关名词。
3. **指标贴合实际**：R 段用可验证的工程/业务指标（延迟、准确率、召回、成本、吞吐、人工耗时等）；量级要像该项目真实能做出的改进，禁止夸张「准确率 99.9%」等空话。
4. **必须有对照**：每个关键指标尽量写成「传统/朴素基线（如规则、单模型、朴素 RAG、无缓存等）→ 改进方案 → 相对 ↑/↓ xx% 或具体前后数值」；面试官能追问实验设置。
5. **方法选型业务驱动**：参考飞书手册（RAG/Agent/LangChain/LangGraph/微调等），按项目业务选合适方案，不要为了新而新；禁止一律 GraphRAG。
6. 输出结构：项目标题 → **技术栈：** → **项目介绍（STAR）** 四段写满；技术名 **加粗**。
"""


def _rag_context(query: str, top_k: int = 6) -> tuple[str, List[dict]]:
    """检索飞书知识库，返回 (format_context, hits)。"""
    try:
        rag = get_rag()
        hits = rag.search((query or "").strip(), top_k=max(1, min(top_k, 12)))
        return rag.format_context(hits), hits
    except Exception:
        return "", []


@app.post("/api/optimize-resume")
async def optimize_resume(req: MatchReq):
    use_rag = bool(getattr(req, "use_rag", True))
    rag_block, rag_hits = ("", [])
    if use_rag:
        rag_block, rag_hits = _rag_context(f"{req.jd}\n{req.resume}\nSTAR 项目 技术栈 指标对比", top_k=6)

    rag_sec = ""
    if rag_block:
        rag_sec = f"""
## 飞书知识库检索（STAR/技术选型参考，按业务选用）
{rag_block}
"""

    prompt = f"""你是专业简历顾问。请同时给出：①优化建议 ②按固定模板写好的「完整新简历」。使用中文 Markdown。

## 当前简历
{req.resume}

## 目标岗位 JD
{req.jd}
{rag_sec}
{RESUME_FORMAT_RULES}

{STAR_PROJECT_RULES}

---

# 一、优化建议

## 🔑 应补充的关键词
（直接从 JD 提取）

## ✏️ 条目改写对比（STAR）
对重要项目经历给出（优化版须含技术栈增强说明 + 基线对照指标）：
> **原文：** ...  
> **优化：** ...

## 📐 结构与投递建议
（2–4 条；务必提醒突出：可实习 6 个月以上、一周内到岗、每周 5 天全勤）

# 二、完整新简历（可直接复制投递）

严格按下列结构输出（字段从简历+JD 推断填充；缺失标【待补充】）：

# 姓名｜求职意向
## 个人信息
（必须含：实习意向——可实习 6 个月及以上；一周内到岗；每周 5 天全勤）
## 教育经历
## 专业技能
## 项目经历
（每个项目：标题 → **技术栈：** → **项目介绍（STAR）**；正文技术名 **加粗**；R 段含相对基线的指标对照）
## 获奖经历
（奖项名｜机构｜时间 + 一句含金量说明；没有则【待补充】）
"""
    meta = {
        "rag": {
            "enabled": use_rag,
            "hit_count": len(rag_hits),
            "sources": [{"source": h["source"], "title": h["title"], "score": h["score"]} for h in rag_hits],
        }
    }
    return stream_resp(prompt, meta=meta)


@app.post("/api/resume-template")
async def resume_template(req: ResumeTemplateReq):
    """生成/定制可编辑的简历 Markdown 模板。"""
    jd_sec = f"\n## 目标岗位 JD\n{req.jd}" if req.jd.strip() else "\n（未提供 JD，生成通用技术岗模板）"
    resume_sec = f"\n## 候选人现有简历（用于推断栏目细项）\n{req.resume}" if req.resume.strip() else ""
    prompt = f"""你是简历模板设计师。请输出一份「可编辑 Markdown 简历模板」，供候选人后续填写或让 AI 按模板生成。

{jd_sec}{resume_sec}

{RESUME_FORMAT_RULES}

## 输出要求
1. 只输出模板正文（Markdown），不要前言后语。
2. 必须包含且按序：个人信息、教育经历、专业技能、项目经历、获奖经历。
3. 个人信息模板中必须预留「实习意向」一行：可实习 6 个月及以上；一周内到岗；每周 5 天全勤。
4. 项目经历至少预留 2 个项目槽位；每个项目先有 **技术栈：** 行，再有 STAR 四段占位。
5. 获奖经历至少预留 1–2 条占位（奖项｜机构｜时间 + 简要说明）。
6. 占位符用 {{{{中文提示}}}} 或简洁示例文字；可结合 JD 把「专业技能」分类写得更贴岗位。
7. 可在模板顶部用 HTML 注释写 3 条使用说明（<!-- ... -->）。

可参考骨架：
{DEFAULT_RESUME_TEMPLATE}
"""
    return stream_resp(prompt)


@app.post("/api/resume-fill")
async def resume_fill(req: ResumeFillReq):
    """按用户导入/编辑的模板，结合 JD+简历自动填写完整简历。"""
    if not req.template.strip():
        return JSONResponse({"error": "请先提供简历模板"}, status_code=400)
    if not req.resume.strip():
        return JSONResponse({"error": "请提供原始简历内容"}, status_code=400)

    use_rag = bool(getattr(req, "use_rag", True))
    rag_block, rag_hits = ("", [])
    if use_rag:
        rag_block, rag_hits = _rag_context(f"{req.jd}\n{req.resume}\n项目 STAR 技术栈 RAG Agent 微调", top_k=6)
    rag_sec = f"\n## 飞书知识库检索（STAR/技术选型参考）\n{rag_block}\n" if rag_block else ""

    prompt = f"""你是资深简历代写顾问。请严格依据「用户简历模板」的结构与标题，把候选人材料改写成一份完整可投递简历。

## 目标岗位 JD
{req.jd}

## 候选人原始简历 / 材料
{req.resume}

## 用户指定的简历模板（必须遵守其章节结构与顺序）
{req.template}
{rag_sec}
{RESUME_FORMAT_RULES}

{STAR_PROJECT_RULES}

## 额外要求
1. 保留模板的一级/二级标题风格；不要擅自改成完全不同的版式。
2. 用真实材料填充；模板里的占位符全部替换掉。
3. 项目经历：每个项目上方 **技术栈：**；STAR 写满；正文技术名 **加粗**；允许基于原项目适当增补技术栈；R 段写基线对照指标。
4. 获奖经历写在项目经历之后；材料无奖项则写【待补充】，禁止编造。
5. 个人信息必须写上：可实习 6 个月及以上；一周内到岗；每周 5 天全勤（除非原材料明确写了冲突信息，再按原材料并标注）。
6. 只输出最终简历 Markdown，不要优化建议、不要代码围栏包裹全文。
"""
    meta = {
        "rag": {
            "enabled": use_rag,
            "hit_count": len(rag_hits),
            "sources": [{"source": h["source"], "title": h["title"], "score": h["score"]} for h in rag_hits],
        }
    }
    return stream_resp(prompt, meta=meta)


@app.post("/api/star-domain-check")
async def star_domain_check(req: StarDomainCheckReq):
    """Agent：检测把项目业务背景换成指定领域是否合理（参考飞书知识库）。"""
    project = (req.project or "").strip()
    domain = (req.domain_bg or "").strip()
    if not project:
        return JSONResponse({"error": "请先粘贴原项目经历"}, status_code=400)
    if not domain:
        return JSONResponse({"error": "请输入目标项目背景（如：法律、医疗、电商）"}, status_code=400)

    use_rag = bool(getattr(req, "use_rag", True))
    rag_block, rag_hits = ("", [])
    if use_rag:
        rag_block, rag_hits = _rag_context(
            f"{project}\n业务背景替换为{domain}\nRAG Agent 场景适配 领域迁移 评测",
            top_k=6,
        )
    rag_sec = f"\n## 飞书知识库检索\n{rag_block}\n" if rag_block else ""

    prompt = f"""你是「项目背景迁移」审核 Agent。判断：把下列项目的业务背景替换为「{domain}」是否合理、面试是否站得住。

## 判定原则
1. 技术主线（模型/系统架构/数据流/工程能力）应能迁移；只换行业话术、数据对象、合规约束、业务指标。
2. 若项目强绑定原行业专有数据/牌照/硬件/不可迁移流程，判 reject。
3. 若可换但需改数据源、评测指标、合规点，判 risky，并给出必须改什么。
4. 若场景与技术天然契合（如 RAG/文档问答→法律；推荐→电商），判 ok。
5. 参考飞书知识库中的 RAG/Agent/评测等通用方法，说明迁移后怎么落地；勿编造手册没有的具体法规条文编号。

## 原项目
{project[:6000]}
{rag_sec}

请只输出一个 JSON 对象（不要 Markdown 围栏），字段如下：
{{
  "verdict": "ok|risky|reject",
  "score": 0到100的整数,
  "summary": "一句话结论",
  "reasons": ["理由1", "理由2"],
  "keep_tech": ["可保留的技术点"],
  "must_change": ["必须改写的业务/数据/指标点"],
  "domain_fit": "迁移后在该领域的典型落地形态（2-4句）",
  "suggestions": "若要继续改写的操作建议"
}}
verdict 含义：ok=可以替换；risky=可换但需谨慎改写；reject=不建议替换。
"""
    llm = get_llm(temperature=0.2, max_tokens=1200)
    try:
        resp = await llm.ainvoke(prompt)
        text = (resp.content or "").strip()
    except Exception as e:
        return JSONResponse({"error": f"检测失败：{e}"}, status_code=502)

    # 抽取 JSON
    data = None
    try:
        data = json.loads(text)
    except Exception:
        m = re.search(r"\{[\s\S]*\}", text)
        if m:
            try:
                data = json.loads(m.group(0))
            except Exception:
                data = None
    if not isinstance(data, dict):
        data = {
            "verdict": "risky",
            "score": 50,
            "summary": "模型未返回标准 JSON，请人工判断后谨慎替换",
            "reasons": [text[:400] or "解析失败"],
            "keep_tech": [],
            "must_change": [],
            "domain_fit": "",
            "suggestions": "可再点一次检测，或换一个更贴近原项目能力的背景词",
        }

    verdict = str(data.get("verdict") or "risky").strip().lower()
    if verdict not in ("ok", "risky", "reject"):
        # 中文兼容
        if "不建议" in verdict or "不可" in verdict:
            verdict = "reject"
        elif "谨慎" in verdict or "风险" in verdict:
            verdict = "risky"
        elif "可以" in verdict or "适合" in verdict:
            verdict = "ok"
        else:
            verdict = "risky"
    data["verdict"] = verdict
    try:
        data["score"] = int(data.get("score") or 0)
    except Exception:
        data["score"] = 0

    data["rag"] = {
        "enabled": use_rag,
        "hit_count": len(rag_hits),
        "sources": [{"source": h["source"], "title": h["title"], "score": h["score"]} for h in rag_hits],
    }
    data["domain_bg"] = domain
    data["raw_preview"] = text[:500]
    return data


@app.post("/api/star-optimize")
async def star_optimize(req: StarOptimizeReq):
    """简历优化栏小项：基于飞书知识库对单个/多个项目做 STAR 专项改写。"""
    project = (req.project or "").strip()
    if not project:
        return JSONResponse({"error": "请粘贴要优化的项目经历（或 STAR 草稿）"}, status_code=400)

    use_rag = bool(getattr(req, "use_rag", True))
    enhance = bool(getattr(req, "enhance_stack", True))
    jd = (req.jd or "").strip()
    resume = (req.resume or "").strip()
    domain = (req.domain_bg or "").strip()
    verdict = (req.domain_verdict or "").strip().lower()
    force_domain = bool(getattr(req, "force_domain", False))

    if domain and verdict == "reject" and not force_domain:
        return JSONResponse(
            {
                "error": f"背景「{domain}」经 Agent 检测为不建议替换。请换背景、先通过检测，或勾选「强制替换」后重试。",
                "verdict": "reject",
            },
            status_code=400,
        )

    rag_q = f"{project}\n{jd}\nSTAR 项目优化 技术栈 基线对比 RAG Agent 微调 评测指标"
    if domain:
        rag_q += f"\n业务背景迁移到{domain} 场景适配"
    rag_block, rag_hits = ("", [])
    if use_rag:
        rag_block, rag_hits = _rag_context(rag_q, top_k=7)

    rag_sec = f"\n## 飞书知识库检索结果（技术选型与写法参考）\n{rag_block}\n" if rag_block else ""
    jd_sec = f"\n## 目标岗位 JD（对齐用）\n{jd}\n" if jd else "\n（未提供 JD，按项目本身与通用大模型岗优化）\n"
    resume_sec = f"\n## 简历上下文（可选）\n{resume[:4000]}\n" if resume else ""
    enhance_line = (
        "允许在原项目合理范围内增补 2～5 个技术点（优先知识库中较新、可落地方法），并在 A 段说明动机与收益。"
        if enhance
        else "不要大幅增补技术栈，以润色表述与补齐 STAR/指标对照为主。"
    )

    if domain:
        domain_sec = f"""
## 项目背景替换（用户指定）
- 目标业务背景：**{domain}**
- Agent 检测结论：{verdict or '未检测（将按可迁移方式谨慎改写）'}
- 要求：
  1. 保留原项目的技术主线与你的贡献角色，不要换成另一个无关系统。
  2. 将 S/T 中的业务场景、数据对象、干系人、合规约束改写为「{domain}」领域真实可信的设定。
  3. A 段技术可保留并按领域微调（如检索语料、评测集、权限）；R 段指标贴合该领域（如审查耗时、召回、误报、人工复核成本等），并与基线方法对照。
  4. 若检测为 risky：优先落实 must_change，避免硬套行业黑话。
  5. 不要编造具体不可核实的牌照/判例编号；用合理业务表述即可。
"""
    else:
        domain_sec = "\n（未指定背景替换：保持原业务背景，只做 STAR/技术栈/指标优化）\n"

    prompt = f"""你是资深简历项目教练 + 工程面试官。请把候选人的「原项目」改写成面试可讲的 STAR 项目经历。

{STAR_PROJECT_RULES}

## 增补策略
{enhance_line}
{domain_sec}
{jd_sec}
{resume_sec}
## 原项目材料（必须基于此改写，不要换成无关项目）
{project}
{rag_sec}

---

请按下列结构输出中文 Markdown（可一次改写材料中的多个项目，每个项目一块）：

# STAR 项目优化结果
{"（已将业务背景迁移为：" + domain + "）" if domain else ""}

（若有多个项目，用 ### 项目 N：标题 分开）

### 项目标题
**技术栈：** ...

**项目介绍（STAR）：**
- **S（情境）：** ...{"（须体现目标背景：" + domain + "）" if domain else ""}
- **T（任务）：** ...
- **A（行动）：** ...（写清关键设计；技术名加粗；增补技术要说明解决什么问题）
- **R（结果）：** ...（含「基线方法 vs 改进方法」对照表或要点，指标贴合实际）

## 背景迁移与增补说明（简短）
- 原背景 → 新背景做了哪些替换
- 相对原文增补了哪些技术栈 / 方法（来自知识库或工程常识）
- 指标对照一览（基线 → 改进 → 变化）

## 面试追问预备（3 题）
简要问答要点即可。
"""
    meta = {
        "rag": {
            "enabled": use_rag,
            "hit_count": len(rag_hits),
            "sources": [{"source": h["source"], "title": h["title"], "score": h["score"]} for h in rag_hits],
        },
        "star_optimize": {
            "enhance_stack": enhance,
            "domain_bg": domain,
            "domain_verdict": verdict,
            "force_domain": force_domain,
        },
    }
    return stream_resp(prompt, temperature=0.4, max_tokens=4096, meta=meta)


@app.post("/api/cover-letter")
async def cover_letter(req: CoverLetterReq):
    prompt = f"""请为{req.name}撰写一封投递{req.company}的专业求职信，使用正式信件格式。

## 简历信息
{req.resume}

## 目标岗位 JD
{req.jd}

---

**写作要求：**
- 开头：明确应聘岗位，体现对公司/岗位的了解
- 中间：选 3 个与 JD 最匹配的核心经历，用数据支撑
- 结尾：表达热情，诚邀面试
- 语气：专业自信，不卑不亢
- 长度：500-600 字

求职信正文："""
    return stream_resp(prompt)

@app.post("/api/interview-qa")
async def interview_qa(req: InterviewReq):
    resume_sec = f"\n## 候选人简历\n{req.resume}" if req.resume.strip() else ""
    prompt = f"""你是面试官，根据以下 JD 生成 10-12 道高频面试题并给出回答框架，使用 Markdown 格式输出中文结果。

## 岗位 JD
{req.jd}{resume_sec}

---

## 一、自我介绍 & 背景

**Q1：** ...  
💡 **回答要点：** ...

## 二、专业技能考察（3-4 题）

**Q2：** ...  
💡 **回答要点：** ...

## 三、行为面试题（STAR 法则，2-3 题）

**Q5：** 描述一次...  
💡 **STAR 框架：**  
- 情境（S）：...  
- 任务（T）：...  
- 行动（A）：...  
- 结果（R）：...

## 四、场景/压力题（1-2 题）

**Q8：** 如果遇到...你会怎么做？  
💡 **回答思路：** ...

## 五、推荐反问问题（面试末尾可以问面试官）
- ...
- ..."""
    return stream_resp(prompt)


@app.get("/api/job-match/meta")
async def job_match_meta():
    """岗位匹配：招聘渠道与大中小厂档位说明。"""
    return {
        "sources": list_sources(),
        "fame_tiers": list_fame_tiers(),
        "companies": list_company_presets(),
        "note": "支持勾选/自填目标公司；岗位链接新标签打开；本页直接展示完整 JD。",
    }


@app.post("/api/job-match")
async def job_match(req: JobMatchReq):
    """根据简历联网检索真实岗位详情页，界面展示 JD；LLM 做匹配排序建议。"""
    resume = (req.resume or "").strip()
    if not resume:
        return JSONResponse({"error": "请先填写或上传简历"}, status_code=400)

    tiers = [t.strip() for t in (req.fame_tiers or []) if t and t.strip()]
    allow = {"large", "mid", "small"}
    tiers = [t for t in tiers if t in allow] or ["large", "mid", "small"]
    tier_labels = {"large": "大厂", "mid": "中厂", "small": "小厂"}
    tier_s = "、".join(tier_labels[t] for t in tiers)

    companies = [c.strip() for c in (req.companies or []) if c and str(c).strip()]
    # 也支持前端误传字符串
    if len(companies) == 1 and ("," in companies[0] or "，" in companies[0]):
        companies = [x.strip() for x in re.split(r"[,，、;；]", companies[0]) if x.strip()]

    if (req.keywords or "").strip():
        kws = [x.strip() for x in req.keywords.split(",") if x.strip()]
    else:
        kws = extract_keywords_heuristic(resume, limit=8)

    city = (req.city or "").strip()
    role = (req.role_hint or "").strip()
    max_hits = max(8, min(int(req.max_hits or 20), 36))

    import asyncio

    loop = asyncio.get_event_loop()
    try:
        hits = await loop.run_in_executor(
            None,
            lambda: search_jobs(
                kws,
                city=city,
                role_hint=role,
                tiers=tiers,
                companies=companies,
                max_hits=max_hits,
            ),
        )
    except Exception as e:
        return JSONResponse({"error": f"联网检索失败：{e}"}, status_code=502)

    jobs_payload = hits_to_dict(hits)
    table = hits_to_markdown(hits)
    sources_s = "、".join(s["name"] for s in list_sources())
    cos_s = "、".join(companies) if companies else "不限"
    resume_snip = resume[:6000]
    prompt = f"""你是资深求职顾问。根据候选人简历，对「联网检索到的真实岗位详情页」做匹配排序与投递建议。

## 硬性要求
1. **只依据下方检索结果中的岗位链接推荐**，不要编造不存在的公司/岗位/URL。
2. 每条结果已是**岗位详情页**，且界面已展示完整 JD；推荐时必须保留原 URL。
3. 厂牌档位按**品牌知名度**：大厂=高知名度头部、中厂=有一定知名度、小厂=知名度较低或成长型；本次用户选择：{tier_s}。
4. 目标公司：{cos_s}（若指定了公司，优先围绕这些公司给建议）。
5. 优先推荐与简历技能/项目匹配度高的岗位；说明匹配点与风险点。
6. 输出中文 Markdown；表格中保留原链接。
7. 若结果偏少，诚实说明，并给出可改进的搜索关键词/公司建议。

## 候选人简历
{resume_snip}

## 检索条件
- 城市偏好：{city or '不限'}
- 意向岗位：{role or '从简历推断'}
- 关键词：{('、'.join(kws) if kws else '（自动抽取）')}
- 目标公司：{cos_s}
- 厂牌档位：{tier_s}
- 覆盖渠道（尽量）：{sources_s}

## 联网检索到的岗位（含 JD 摘要）
{table}

---

请按下列结构输出（界面已单独展示岗位卡片与 JD，这里侧重排序与建议）：

## 1. 简历速读（意向方向 / 核心技能 / 城市推断，各 1～2 句）

## 2. Top 推荐（最多 8 条，必须用上方真实 URL）
| 排名 | 匹配度 | 档位 | 公司 | 岗位 | 为何适合 | 链接 |
|:----:|:------:|:----:|------|------|----------|------|
| … | 高/中/低 | 大厂/中厂/小厂 | … | … | … | URL |

## 3. 投递策略
- 本周优先顺序
- 简历需补强的 2～3 点

## 4. 下次检索建议（关键词 / 城市 / 公司 / 档位）
"""
    meta = {
        "jobs": jobs_payload,
        "job_match": {
            "keywords": kws,
            "city": city,
            "role_hint": role,
            "companies": companies,
            "fame_tiers": tiers,
            "hit_count": len(hits),
            "sources": [s["name"] for s in list_sources()],
            "hint": "岗位链接请在新标签页打开；本页已展示完整 JD",
        },
    }
    return stream_resp(prompt, temperature=0.35, max_tokens=4096, meta=meta)


def _resolve_deploy_flags(include_github: bool, include_demo: bool, include_deploy: Optional[bool]):
    """拆分 GitHub / 在线 Demo；旧字段 include_deploy 同时控制两者。"""
    if include_deploy is not None:
        return bool(include_deploy), bool(include_deploy)
    return bool(include_github), bool(include_demo)


def _normalize_project_plan(count: int = 3, tiers: Optional[List[str]] = None) -> List[dict]:
    """规范化「生成几个 + 质量定位」。"""
    order = ["主推", "备选", "加分"]
    alias = {
        "主推": "主推", "备选": "备选", "加分": "加分",
        "main": "主推", "backup": "备选", "bonus": "加分",
        "primary": "主推", "alt": "备选", "extra": "加分",
    }
    try:
        n = int(count or 3)
    except (TypeError, ValueError):
        n = 3
    n = max(1, min(3, n))
    cleaned = []
    for t in (tiers or []):
        key = str(t).strip()
        mapped = alias.get(key) or alias.get(key.lower()) or (key if key in order else None)
        if mapped and mapped not in cleaned:
            cleaned.append(mapped)
    if len(cleaned) < n:
        for t in order:
            if t not in cleaned:
                cleaned.append(t)
            if len(cleaned) >= n:
                break
    cleaned = cleaned[:n]
    return [
        {
            "idx": i + 1,
            "tier": t,
            "tag": f"【{t}】",
            "title": f"项目 {i + 1}：【{t}】",
        }
        for i, t in enumerate(cleaned)
    ]


# 飞书手册提炼：全链路较新/有效技术选型（按项目业务背景选用，禁止一律硬塞）
_FEISHU_TECH_ROUTE = """
## 飞书知识库技术路线参考（全链路按业务选型）
**总原则**：先写清本项目业务对象、数据形态、交互形态、约束（时延/成本/准确率），再从手册里选「较新 + 效果更好」的方法；**禁止**无理由全员 GraphRAG / 全员 Supervisor / 全员微调。能力覆盖仍须齐全，但**具体实现手法按业务选型**，并在覆盖表写「为何选 A 不选 B」。

### 1) RAG（《RAG手册》等作参考池，非必选菜单）
- **选型逻辑**：先根据业务查询形态（FAQ / 长文档 / 多跳关系 / 需要纠错 / 低延迟等）与约束，选**最合适且效果更好**的检索增强方案；相对基线「单路向量 TopK」须写清为何更好、如何验收。
- **手册方法仅作可选参考示例**（可选用、可组合、也可选手册外更贴业务的有效方案，**禁止**把 Hybrid+RRF / RAPTOR / CRAG 写成必选项）：语义分块、Parent Document、Multi-representation、RAPTOR、ColBERT、Multi-Query、RAG-Fusion、RRF、HyDE、Decomposition、Step-Back、Adaptive Routing、Hybrid、Rerank、CRAG、Agentic RAG、GraphRAG 等。
- **GraphRAG**：仅当业务确需实体关系/多跳/社区级总结时再选；否则不要写。
- 简单场景用轻量有效方案即可；复杂场景再上重链路——以效果与可落地为准。

### 2) Agent 编排（《Agent手册》《Langgraph手册》《面试题》）
- 单 Agent：探索任务用 **ReAct**；步骤清晰用 **Plan-and-Execute**；需自检可用 Reflection。
- Multi-Agent：业务真需要多角色时再拆——**Supervisor** 调度 / **Handoff** 交接 / 评测·审查 Agent；禁止无职责堆砌。
- LangGraph：**StateGraph** + **checkpointer**（会话/断点续跑）；敏感步骤 **Human-in-the-loop / interrupt**。
- 每个 Agent 必须有对应 **skill**（映射表），失败回退写清。

### 3) 工具与 MCP（《Agent》《langgraph+mcp 项目》《Harness》）
- Function Calling：严格 schema、校验、错误回灌；相对「模型空想」。
- MCP：STDIO / HTTP·SSE；多 Server 组合；工具可观测；长任务与补参按业务需要。
- Harness 思维：权限边界、测试夹具、日志留痕、完成标准=可验证产物（不是「模型说做完了」）。

### 4) 记忆（《Agent手册》）
- 短期：窗口 / **摘要记忆**；长期：向量/结构化记忆 + 重要性与时效清理。
- 按业务选：客服多轮偏短期摘要；个性化推荐偏长期画像；一次性工具任务可轻量记忆。

### 5) 推理与生成（《Transformer》《大模型八股》《Harness》）
- 长上下文：RoPE 缩放 / YaRN / 滑动窗口——相对硬截断。
- 服务侧（贴业务）：vLLM/SGLang、投机解码、KV Cache、批推理——须写清延迟/成本收益。
- 流式：SSE + 三态 UI；TTFT 可测。

### 6) 后训练与工程落地（《大模型八股》《Vibe Coding》《Codex》《Claude Code》）
- 微调开关开：SFT/LoRA/QLoRA（可选 DPO），小数据可落地；关：用已后训练权重 + 推理侧策略。
- 工程：可复现脚本、PIPELINE 逐步通关、案例+真实数据双通道评测、国内可访问 Demo（按开关）。

### 硬性比例（不变）
- 技术栈 **≥8**；**较新栈 ≥3**（写手册里的具体新点+业务理由）；**自研改进 ≥2**（基线→做法→↑/↓%→测法）。
"""


@app.post("/api/project-recommend")
async def project_recommend(req: ProjectRecReq):
    """根据 JD 推荐 3 个可落地项目；微调 / GitHub / Demo 可分别开关。"""
    bg = f"\n## 候选人背景（可选）\n{req.background}" if req.background.strip() else ""
    pref = f"\n## 技术/工期偏好（可选）\n{req.preference}" if req.preference.strip() else ""
    env_parts = []
    if req.python_version.strip():
        env_parts.append(f"- Python：{req.python_version.strip()}")
    if req.cuda_version.strip():
        env_parts.append(f"- CUDA：{req.cuda_version.strip()}")
    if req.gpu_model.strip():
        env_parts.append(f"- GPU：{req.gpu_model.strip()}")
    env = ("\n## 本机训练/推理环境（必须写入方案，不可超出）\n" + "\n".join(env_parts)) if env_parts else ""

    do_ft = bool(req.include_finetune)
    do_pt = bool(getattr(req, "include_pretrain", False))
    do_iv = bool(getattr(req, "include_interview_qa", False))
    do_mm = bool(getattr(req, "include_multimodal", False))
    do_diff = bool(getattr(req, "differentiate", True))
    do_gh, do_demo = _resolve_deploy_flags(req.include_github, req.include_demo, req.include_deploy)
    do_dp = do_gh or do_demo

    region = (req.llm_region or "cn").strip().lower()
    if region not in ("cn", "overseas", "local"):
        region = "cn"
    region_label = {"cn": "国内大模型 API", "overseas": "国外大模型 API", "local": "纯本地开源（不调用云 API）"}[region]
    provider = (req.llm_provider or "").strip() or "（未指定厂商，请按路线推荐常见选项）"
    model = (req.llm_model or "").strip() or "（未指定模型，请按厂商给默认推荐）"
    llm_scope = (
        f"\n## 推理大模型 API 选型（必须写入每个项目的技术栈、STAR 与 AI 搭建提示词）\n"
        f"- 路线：{region_label}\n"
        f"- 厂商/平台：{provider}\n"
        f"- 具体模型：{model}\n"
        f"- 硬性：默认主路径必须使用上述选型；若微调开启，基座/旁路也要与该路线一致或说明关系。\n"
        f"- 国内路线：写国内可直连的官方/兼容 Base URL，禁止默认写成必须翻墙才能调的 OpenAI 官方。\n"
        f"- 国外路线：写清厂商与模型；可附一句网络前提；可给可选的国内兼容中转，但主方案仍按用户所选国外厂商。\n"
        f"- 纯本地：禁止依赖云 API Key；写 Ollama/vLLM/llama.cpp 等本地推理与模型权重来源。\n"
    )

    scope = (
        f"\n## 交付范围开关（必须严格遵守）\n"
        f"- 是否包含模型微调（后训练）：{'是（必须写微调/后训练方案并写入 AI 提示词）' if do_ft else '否（禁止要求微调/LoRA/训练流水线；可用 API 推理即可）'}\n"
        f"- 是否包含预训练/领域 CPT：{'是（可写小规模继续预训练 CPT；禁止从零预训练）' if do_pt else '否（不要写 CPT/继续预训练专章；直接复用开源预训练权重即可，一句带过）'}\n"
        f"- 是否包含面试相关问题：{'是（每个项目必须写「面试追问/相关问题」专章：题目+简要回答框架）' if do_iv else '否（禁止输出面试追问/模拟面试 Q&A）'}\n"
        f"- 是否支持多模态：{'是（项目须含图像/音视频等非纯文本模态的输入或理解链路，写清模型与数据）' if do_mm else '否（默认纯文本/结构化文本链路即可，不要硬塞多模态）'}\n"
        f"- 是否 GitHub 公开仓：{'是（必须写仓库结构 + README + 推送步骤）' if do_gh else '否（不要强制推仓；本地可运行即可）'}\n"
        f"- 是否在线 Demo：{'是（必须国内可访问公网 Demo）' if do_demo else '否（禁止要求上线公网 Demo）'}\n"
    )

    detail_raw = (req.detail_level or "detailed").strip().lower()
    is_brief = detail_raw in ("brief", "short", "简短", "简短版")
    detail_label = "简短版" if is_brief else "详细版"
    plan = _normalize_project_plan(req.project_count, req.project_tiers)
    proj_n = len(plan)
    tier_names = "、".join(p["tier"] for p in plan)
    try:
        days = int(req.timeline_days or 7)
    except (TypeError, ValueError):
        days = 7
    if days not in (1, 3, 7, 15, 30):
        # 就近归一
        days = min((1, 3, 7, 15, 30), key=lambda x: abs(x - days))
    timeline_label = f"{days} 天"
    scale_raw = (req.project_scale or "school").strip().lower()
    if scale_raw in ("enterprise", "企业", "企业级", "企业级项目", "company", "corp", "industry"):
        is_enterprise = True
    else:
        is_enterprise = False  # school / 学校级 等默认学校级
    scale_label = "企业级项目" if is_enterprise else "学校级项目"
    if is_enterprise:
        scale_scope = (
            f"\n## 项目背景（必须严格遵守）：{scale_label}\n"
            f"- 叙事定位：实习生 / 初级工程师在**真实业务链路**中的可落地项目（仍**禁止任何公司名/品牌名**）。\n"
            f"- Situation 用「业务对象 + 流程环节 + 生产约束」：峰值流量、SLA、脏数据、多角色协作、可观测/回滚等。\n"
            f"- 约束侧重：线上可用性、成本、延迟、并发、权限与审计、灰度发布（按 JD 选型，勿硬塞）。\n"
            f"- Result 指标优先：线上/准线上可验证的延迟、吞吐、准确率、成本、故障恢复等。\n"
            f"- 角色口吻：对业务方/协作方交付，而不是课程作业答辩。\n"
        )
        scale_biz_rule = (
            "5. **业务背景硬性要求（企业级项目；以实际业务为主，禁止公司名）**：\n"
            "   - **禁止**出现任何真实/虚构公司名、品牌名、产品商品名（含「某大厂」「字节/阿里/腾讯」等指代）。\n"
            "   - **禁止**照抄 JD 业务原文；只抽象「业务环节 + 痛点」。\n"
            "   - 用「业务对象 + 流程环节」描述，例如：日活内容审核队列、实验指标回流链路、多源日志告警收敛、离线评测批跑、候选召回后重排等。\n"
            "   - 每个项目必须先写「痛点清单」（≥3 条：数据质量/延迟/成本/准确率/稳定性/协作），再映射到功能设计。\n"
            "   - 必须有真实约束 ≥2：时限、成本、脏数据、并发、可观测性、回滚等。\n"
            "   - STAR 的 S 必须像**企业业务场景**（生产/准生产），不要写成课程作业或毕设开题。"
        )
    else:
        scale_scope = (
            f"\n## 项目背景（必须严格遵守）：{scale_label}\n"
            f"- 叙事定位：课程设计 / 毕业设计 / 实验室课题 / 学生团队作品集（**禁止编造学校全名以外的「伪企业」外壳**；也**禁止真实公司名**）。\n"
            f"- Situation 用「课题目标 + 数据/算力约束 + 验收场景」：开源/公开数据集、本机或校园算力、答辩/课程验收、组员分工。\n"
            f"- 约束侧重：工期短、显存/预算有限、可复现实验、文档与演示清晰；可写导师/课程评分维度，但不要写成假实习证明。\n"
            f"- Result 指标优先：离线评测、消融对比、Demo 验收、相对基线的提升；可附一句「若上线可如何验证」。\n"
            f"- 角色口吻：学生负责人 / 核心开发，强调独立完成与可讲解深度；技术深度仍须对齐 JD。\n"
        )
        scale_biz_rule = (
            "5. **业务背景硬性要求（学校级项目；痛点驱动，禁止公司名）**：\n"
            "   - **禁止**出现任何真实/虚构公司名、品牌名；也不要伪装成「某互联网大厂实习项目」。\n"
            "   - 场景写成**课程/毕设/实验室/学生团队**可解释的课题：业务对象可来自公开领域问题（内容审核队列、日志异常定位、检索评测等），但交代是课题/作品集而非在职交付。\n"
            "   - 每个项目必须先写「痛点清单」（≥3 条：数据质量/延迟/成本/准确率/可复现/协作），再映射到功能设计。\n"
            "   - 必须有真实约束 ≥2：算力/显存、工期、公开数据噪声、可复现、答辩演示稳定性等。\n"
            "   - STAR 的 S 必须像**学校级课题背景**，T/A/R 仍要有工程深度与可验证指标（对齐 JD 技能）。"
        )

    form_raw = (req.project_form or "web").strip().lower()
    if form_raw in ("extension", "plugin", "browser", "浏览器插件", "插件", "chrome", "edge"):
        form_key = "extension"
        form_label = "浏览器插件"
    elif form_raw in ("desktop", "pc", "电脑端", "桌面", "桌面端", "desktop-app", "electron", "tauri"):
        form_key = "desktop"
        form_label = "电脑端 App"
    elif form_raw in ("mobile", "app", "移动端", "移动应用", "ios", "android", "手机端"):
        form_key = "mobile"
        form_label = "移动端 App"
    else:
        form_key = "web"
        form_label = "网站"
    if form_key == "extension":
        form_scope = (
            f"\n## 项目形式（必须严格遵守）：{form_label}\n"
            f"- 交付形态：Chrome / Edge 等 **Manifest V3 浏览器扩展**（可含 popup / side panel / content script / background service worker）。\n"
            f"- 技术选型优先：扩展前端（HTML/JS/TS/React 等）+ 可选本地/远端后端 API；权限最小化；说明 content script 与页面交互边界。\n"
            f"- 演示：本地加载未打包扩展 +（若开 Demo）国内可访问的说明页/录屏或配套站点；禁止只做纯网页却声称是插件。\n"
            f"- 一览表与定位句必须标明「浏览器插件」；AI 提示词须含扩展目录结构、manifest、安装/加载步骤。\n"
            f"- 评测可在目标网页上挂载跑案例；真实数据可通过 popup 上传或与配套 API 导入。\n"
        )
    elif form_key == "desktop":
        form_scope = (
            f"\n## 项目形式（必须严格遵守）：{form_label}\n"
            f"- 交付形态：**电脑端桌面 App**（优先 Electron / Tauri / Flutter Desktop，或按 JD 写原生；须在方案里写死选型）。\n"
            f"- 含：窗口/托盘/本地文件权限、本地或远端 API、三态 UI；说明与纯网页的差异（离线能力、系统集成等，按选题选型）。\n"
            f"- 演示：本机安装包或开发模式一键启动可打开主窗口 +（若开 Demo）国内可访问的下载页/录屏说明；禁止做成只有浏览器网页却声称是桌面 App。\n"
            f"- 一览表与定位句必须标明「电脑端 App」；AI 提示词须含工程结构、打包/启动命令、Windows 验收步骤。\n"
            f"- 评测：桌面端「跑案例」入口 + 本地文件导入真实数据。\n"
        )
    elif form_key == "mobile":
        form_scope = (
            f"\n## 项目形式（必须严格遵守）：{form_label}\n"
            f"- 交付形态：**移动端 App**（优先 Flutter / React Native / uni-app，或按 JD 写原生 iOS/Android；须在方案里写死选型）。\n"
            f"- 含：关键页面导航、本地/远端 API、三态 UI；可附模拟器运行说明。\n"
            f"- 演示：模拟器或真机可跑 MVP +（若开 Demo）国内可访问的安装包/体验页说明；禁止做成只有网页却声称是移动端 App。\n"
            f"- 一览表与定位句必须标明「移动端 App」；AI 提示词须含工程结构、启动命令、真机/模拟器验收。\n"
            f"- 评测：App 内「跑案例」入口 + 导入真实数据（相册/文件选择/粘贴等，按场景选型）。\n"
        )
    else:
        form_scope = (
            f"\n## 项目形式（必须严格遵守）：{form_label}\n"
            f"- 交付形态：**Web 网站**（前后端或可运行的单页+API；浏览器访问主路径）。\n"
            f"- 技术选型优先：FastAPI/Flask 等后端 + 前端单页，或 Next 等全栈（按 JD）；须写清路由与三态 UI。\n"
            f"- 演示：本地一键启动浏览器可打开 +（若开 Demo）国内可访问公网站点。\n"
            f"- 一览表与定位句必须标明「网站」；AI 提示词须含目录、启动命令、主路径 URL。\n"
            f"- 评测：页面上「跑案例」与「上传/导入真实数据」双入口。\n"
        )

    plan_scope = (
        f"\n## 项目数量与质量定位（必须严格遵守）\n"
        f"- 恰好生成 **{proj_n}** 个项目，质量定位依次为：{tier_names}\n"
        f"- 标题必须使用：{'；'.join(p['title'] for p in plan)}\n"
        f"- 不要多写未选中的质量档；不要把未选档硬塞进来。\n"
        f"\n## 项目完成时间（必须严格遵守）\n"
        f"- 每个项目的建议工期 / MVP 交付周期必须按 **{timeline_label}** 规划（可写日内里程碑，但总工期不超过 {days} 天）。\n"
        f"- TODO 与 AI 搭建提示词中的任务量、范围必须匹配 {timeline_label}，禁止按默认「1–2 周」随意拉长。\n"
        f"- 一览表「建议工期」列统一写：{timeline_label}（或更短的可交付切片，但不要超过）。\n"
        f"{scale_scope}{form_scope}"
    )
    _len_brief = (
        f"- 简短版：保留硬性要素（{proj_n} 项目、≥8 技术栈、LLM 全链路覆盖表、完整 STAR、每项目 1 份搭建提示词"
        + ("、面试相关问题" if do_iv else "")
        + "）；删减展开论述。\n"
    )
    _len_full = (
        f"- 详细版：按模板展开必要小节，但仍受每项目 ≤15000 字约束；每项目恰好 1 份搭建提示词"
        + ("；面试开关开启则另写面试相关问题专章" if do_iv else "")
        + "。\n"
    )
    _len_ban = (
        "- **禁止输出**：作品集包装专章、作品集附录、额外叙事附录、第二份及以上提示词；面试相关问题写在正文专章，**不要**写进搭建提示词。\n"
        if do_iv
        else "- **禁止输出**：面试追问 / 模拟面试 Q&A、作品集包装专章、作品集附录、额外叙事附录、第二份及以上提示词。\n"
    )
    length_scope = (
        f"\n## 篇幅与输出范围（必须严格遵守）\n"
        f"- 当前篇幅：{detail_label}\n"
        f"- **字数硬上限：每个项目整体（从该项目标题起，到下一项目标题前）≤ 15000 字**；超了就删水分，禁止灌水。\n"
        + (_len_brief if is_brief else _len_full)
        + "- **每个项目有且仅有 1 个**「交给其他 AI 的自动化搭建提示词」```text 代码块（只用于搭建，不要其它类型提示词）。\n"
        + _len_ban
    )

    must_bits = [
        "对齐 JD 技术能力（技能对齐，不是照搬公司业务）",
        "以真实业务痛点驱动、无公司名的场景",
        "**指标对比验证**（基线 vs 改进，须写清↑/↓百分比）",
        "**每项目≥8 个核心技术栈**",
        "**主流大模型全链路覆盖**（LLM / RoPE / **RAG按业务选型** / Multi-Agent / Function Calling / MCP / 记忆系统 / 流式输出 / 模型部署 / 后训练；预训练按开关；**禁止一律 GraphRAG**）",
        "**技术栈按真实工程流程串联**（不是关键词堆砌）",
        "**设计 Agent 时必须搭配对应 skill**（每个 Agent 写清 skill 名称、职责、工具/MCP、输入输出与验收）",
        "**自研改进硬性 ≥2 个**（每个含传统基线→做法→↑/↓%→测法）",
        "**较新栈≥3=飞书手册方向里的具体新进展**（RAG/Agent/MCP/记忆/推理部署等须写具体新点+业务选型理由，禁止只写空名词、禁止一律硬塞）",
        "**评测双通道：内置案例数据 + 可导入真实数据**",
    ]
    if do_diff:
        must_bits.append("**选题差异化（尽量避免与市面常见 Demo 同质化）**")
    if do_iv:
        must_bits.append("**每项目含面试相关问题专章（题目+简要回答框架）**")
    if do_mm:
        must_bits.append("**支持多模态（图像/音视频等非纯文本链路须落地）**")

    if do_ft:
        must_bits.append("**模型微调方案**")
    if do_gh:
        must_bits.append("面试官能看的 **GitHub 公开仓**")
    if do_demo:
        must_bits.append("**国内可访问在线 Demo**")
    must_line = "、".join(must_bits)

    rules = [
        f"1. 只推荐 **{proj_n}** 个项目，质量定位依次为：{tier_names}；彼此方向必须明显不同。",
        "2. 技术栈从 JD 抽取并对齐，写清选型理由与「明确不做」；禁止无关堆料、禁止只有 CRUD 老三样。",
        (
            "3. **选题差异化（已开启）**：\n"
            "   - 尽量避免与市面常见 Demo 同质化；优先做有清晰业务对象与流程的系统，例如：风控审核流水线、实验/评测平台、内容生产质检、检索重排评测台、推理成本治理、多模态质检、推荐重排、日志异常定位、数据标注提效、仿真评测、调度优化等（按 JD 技能选型，勿硬套）。\n"
            "   - 若选做智能客服 / 问答助手 / 知识库问答等常规方向，也**必须**写清「差异化卖点」：相对市面常见 Demo 你多解决了什么业务难点、多了哪些工程/算法改进。\n"
            "   每个项目必须写清「差异化卖点」一句话。"
            if do_diff
            else
            "3. **选题差异化（已关闭）**：\n"
            "   - 不强制差异化，可直接做智能客服 / 客服机器人 / 通用问答助手 / 知识库问答 / 企业 FAQ / Chat with PDF 等常规方向。\n"
            "   - 仍要求技术栈按真实工程流程串联、指标可对比；「差异化卖点」可写「无（常规方向）」或简述基础价值即可。"
        ),
        "4. **核心技术栈数量硬性要求（每个项目）**：至少列出 **8 个**核心技术/方法（写成清单或表格），其中：\n"
        "   - **较新栈（≥3 项，硬性）含义（重要）**：从飞书手册（RAG/Agent/LangGraph/MCP/Harness/Transformer/八股/工程实战）里选**较新且效果更好**的具体进展，并落到本项目；**必须先写业务背景再选型**，禁止无理由一律 GraphRAG / 一律 Supervisor / 一律微调。优先方向示例（贴业务选，勿全堆）：\n"
        "     · **RAG**：按业务查询形态选合适且效果更好的方案（手册方法仅参考示例，**不强制** Hybrid+RRF / RAPTOR / CRAG / GraphRAG 中的任何一项）；须写清相对朴素向量 TopK 为何更好、如何测。\n"
        "     · **Agent**：ReAct vs Plan-and-Execute（按任务是否探索性）；Supervisor/Handoff；LangGraph checkpointer / Human-in-the-loop；Reflection。\n"
        "     · **工具/MCP/Harness**：严格 Function Calling schema；MCP STDIO 或 HTTP·SSE；权限边界+可验证完成标准（Harness）。\n"
        "     · **记忆**：摘要短期记忆、长期向量记忆、重要性/时效清理——按多轮/个性化需要选型。\n"
        "     · **推理部署**：RoPE/YaRN、vLLM/SGLang、投机解码、KV Cache、SSE 流式——写清延迟/成本「新在哪里」。\n"
        "   - **自研改进（≥2 项，硬性，不可省略）**：自己的方法/工程创新（非调包默认），相对朴素基线可测。可落在检索融合、Agent 握手、记忆策略、路由阈值、缓存批推理、质检门槛等（**不必**强行图谱剪枝）。\n"
        "   - 每个自研改进必须写清：**传统/朴素基线（数值）→ 你的做法 → 改进后（数值）→ ↑/↓ 百分比 → 测法（案例集+真实数据）**。\n"
        "   - 「改进对照表」至少 **2 行自研改进**；技术栈表里类型=「自研改进」的条目也至少 2 条。\n"
        "   - 输出固定小节「## 1.3b 核心技术栈（≥8）+ 自研改进」「## 1.3c LLM 主流技术覆盖表（按工程流程）」与「改进对照表」（收益列必须含 %）。",
        "4b. **主流大模型技术完整性（每个项目必须基本覆盖，且按真实工程流程串联，禁止只堆名词）**：\n"
        "   必须用「覆盖表」逐项写清：技术点 | 在本项目流程中的阶段 | 具体做什么 | 相对传统做法的改进（含↑/↓% 或明确「本开关关闭时的等价落地」）。\n"
        "   **必覆盖技术点（名称可微调，但能力不可缺；预训练按开关）**：\n"
        "   1) **LLM**：推理主路径（严格按用户 API 选型）；写清调用位置与失败重试。\n"
        "   2) **RoPE / 长上下文**：说明位置编码/上下文长度策略（如 RoPE 缩放、YaRN、分块/滑动窗口、超长文档切分策略）；相对「截断硬切」的改进。\n"
        "   3) **RAG（按业务选型，必填）**：先写业务查询形态与约束，再选**合适且效果更好**的检索增强主路径（可参考飞书《RAG手册》方法池，也可选其它贴业务的有效方案）。**禁止**无理由一律 GraphRAG；也**禁止**把 Hybrid+RRF / RAPTOR / CRAG 等写成固定必选。覆盖表须写：选了什么、为何适合本业务、相对「单路向量 TopK」的改进与测法。\n"
        "   4) **Multi-Agent**：≥2 个角色分工（规划/执行/评测/审查等按业务选型）；相对「单 Agent 一把梭」的改进。\n"
        "   5) **Function Calling**：结构化工具调用（参数 schema、校验、错误回灌）；相对「模型空想直接答」的改进。\n"
        "   6) **MCP**：至少 1 个 MCP Server/工具面（本地工具、检索、评测、文件等）；说明与 Function Calling 的协作边界。\n"
        "   7) **记忆系统**：短期（会话/工作记忆）+ 长期（摘要/向量/结构化记忆）至少一种可落地设计；相对「无状态每轮重传」的改进。\n"
        "   8) **流式输出**：SSE/流式 token 推送到前端；相对「整段阻塞返回」的改进（TTFT/体感延迟）。\n"
        "   9) **模型部署**：可演示的推理服务形态（API / vLLM·Ollama·本地服务 / 国内可访问 Demo，按开关）；写清启动与验收。\n"
        + (
            "   10) **预训练（已开启）**：可写小规模领域继续预训练 CPT；禁止从零预训练；写清数据/算力/验收。\n"
            if do_pt
            else "   10) **预训练（已关闭）**：不要写 CPT/继续预训练专章；一句说明直接复用开源预训练 checkpoint 即可。\n"
        )
        + "   11) **后训练**：若微调开关=开 → 必须写 SFT/LoRA/QLoRA（可选 DPO）流水线；若=关 → 写「使用已后训练权重 + 推理侧对齐/偏好策略」，禁止假装做了训练。\n"
        "   **流程串联硬性要求**：技术栈必须映射到端到端阶段顺序，推荐骨架为：\n"
        + (
            "   数据/知识构建 → 预训练/CPT → 后训练或选用后训练权重 → **RAG（按业务选型）** → Multi-Agent → Function Calling/MCP → 记忆 → LLM推理(RoPE) → 流式 → 部署 → 评测。\n"
            if do_pt
            else "   数据/知识构建 →（跳过 CPT）选用预训练权重 → 后训练或选用后训练权重 → **RAG（按业务选型）** → Multi-Agent → Function Calling/MCP → 记忆 → LLM推理(RoPE) → 流式 → 部署 → 评测。\n"
        )
        + "   禁止：把上述技术写成互不相关的名词列表；写了「用了 RAG/Multi-Agent/MCP」就必须真正落地对应能力，不能只堆名词；**禁止无理由全员 GraphRAG**。",
        "4b2. **RAG 选型（硬性：业务优先、效果优先；手册仅参考）**：\n"
        "   - 先写 1～2 句业务查询形态与约束，再选型；目标是本业务下检索/生成效果更好、可验收。\n"
        "   - 相对基线「朴素单路向量 TopK」写清机制差异与指标（召回/nDCG/幻觉率/延迟等按业务选）。\n"
        "   - 飞书手册中的 Hybrid/RRF、RAPTOR、CRAG、HyDE、Rerank、GraphRAG 等只是**可选参考**，用哪一种（或组合、或其它有效方案）由业务决定，**无一强制**。\n"
        "   - 不需要图谱就不要写 GraphRAG；不需要重型纠错/层级索引就不要硬上 CRAG/RAPTOR。",
        "4b3. **其它全链路也按飞书手册 + 业务选型（硬性）**：\n"
        "   - **Agent**：按任务选 ReAct / Plan-and-Execute / Supervisor；写清「业务为何需要多 Agent（或不需要）」；LangGraph 状态与 checkpointer/HITL 按是否长任务/敏感操作选型。\n"
        "   - **工具/MCP**：工具面贴业务对象（检索/文件/评测/外部 API）；协议 STDIO vs HTTP·SSE 按部署选；Harness：权限+日志+可验证验收。\n"
        "   - **记忆 / 流式 / 部署**：多轮才加重记忆；交互产品必须 SSE；推理加速（vLLM/投机解码等）仅在延迟/吞吐是痛点时作为较新栈写出收益。\n"
        "   - 覆盖表每一行都要有「传统做法 → 你选的手册方法 → 业务理由」；禁止全项目复制同一套名词清单。",
        "4c. **Agent-Skill 搭配（每个项目硬性，不可省略）**：\n"
        "   - 只要写到 Multi-Agent / Agent 编排，就必须输出固定小节「## 1.3d Agent-Skill 映射表」。\n"
        "   - 每个 Agent 角色都要绑定 1～2 个对应 skill，skill 必须是可实现的能力模块，而不是空泛标签；例如：规划 Agent→task_planning_skill，检索 Agent→hybrid_retrieval_skill / rerank_skill / graph_retrieval_skill（按实际 RAG 选型命名），工具 Agent→mcp_tool_calling_skill，评测 Agent→eval_metric_skill，记忆 Agent→memory_update_skill，部署 Agent→deploy_release_skill。\n"
        "   - 映射表列必须包含：Agent 角色 | 对应 skill | 触发时机 | 输入 | 输出 | 调用工具/MCP/函数 | 失败回退 | 验收指标。\n"
        "   - 技术栈表、LLM 覆盖表、AI 搭建提示词和 TODO 中都要同步出现该映射；禁止只写「有规划 Agent/执行 Agent」但不写 skill。",
        scale_biz_rule,
        "6. **指标对比是验收核心（Result 必须能量化证明「更好」）**：\n"
        "   - ≥3 个可验证指标；至少 1 个性能/轻量化，至少 1 个对应痛点缓解。\n"
        "   - 每个关键指标必须写成「基线值 → 改进后值 → **变化百分比（↑xx% 或 ↓xx%）**」；禁止只写「明显提升」「效果更好」。\n"
        "   - STAR 的 R 必须用这些百分比讲清：因为做了哪项设计（对应 A），所以某指标 ↑/↓ 多少。\n"
        "   - 未实测前可写「目标百分比 + 测量方法」；实测后填真实数；禁止编造无法复现的线上大盘数字。\n"
        "6b. **评测数据双通道（硬性）**：\n"
        "   - **案例数据**：仓库内置 `data/samples/`（或等价）固定样例，一键跑通基线 vs 改进对照，面试可演示。\n"
        "   - **真实数据**：必须提供上传/导入入口（UI 文件上传、或 API/CLI 指定路径），支持用户自有 CSV/JSON/图片/日志等（按项目选型）；导入后走同一套评测流水线并输出对照表。\n"
        "   - 文档与 AI 提示词须写清：样例怎么跑、真实数据格式/字段、导入命令或页面操作、隐私注意（本地处理、不上传第三方）。",
        "7. **若提供了本机环境（Python / CUDA / GPU）**：相关命令必须按该环境可跑；显存不够则降级；禁止默认假设 A100。",
        (
            f"8. **篇幅=简短版，但 {proj_n} 个项目 STAR 仍须完整（可短不可缺）**：\n"
            "   - 每个项目独立输出「面试可讲的完整 STAR」：S / T / A / R 四段都有实质内容。\n"
            "   - **禁止**：「同项目1」「结构同上」「略」「见上文」。\n"
            "   - 字数下限：S≥40字、T≥30字、A≥3条、R≥2条且须含 ↑/↓ 百分比。\n"
            ""
            + (
                "   - 每项目结构精简为：定位+痛点3条+技术栈≥8+LLM覆盖表+改进2条+STAR+**恰好 1 份**搭建提示词+**面试相关问题（3～5题）**；禁止作品集专章。"
                if do_iv
                else "   - 每项目结构精简为：定位+痛点3条+技术栈≥8+LLM覆盖表（可缩成紧凑表）+改进2条（含%相对传统）+（开启项简述）+STAR+**恰好 1 份**搭建提示词；禁止面试追问/作品集专章。"
            )
            + ""
            if is_brief
            else
            f"8. **{proj_n} 个项目的 STAR 必须各自完整写满，禁止略写**：\n"
            "   - 每个项目独立输出完整 STAR：S / T / A / R 四段都要有实质段落。\n"
            "   - **禁止**：一句话带过、只写「S/T/A/R」标签、写「同项目1」「结构同上」「略」「见上文」。\n"
            "   - S≥80字、T≥60字、A≥5条或≥150字（必须含 LLM 全链路技术栈与相对传统的改进点）、R 必须用「设计→指标↑/↓百分比」讲清验证结论。\n"
            ""
            + (
                "   - **必须**另写「面试相关问题」专章（6～8 题，每题含简要回答框架）；不要写进搭建提示词。"
                if do_iv
                else "   - **禁止**另写「面试追问」小节；项目内 STAR 即可，不要追问清单。"
            )
            + ""
        ),
        "9. **推理 API 必须按用户选型落地**：路线=" + region_label + f"；厂商={provider}；模型={model}。"
        " 不得擅自改成另一路线的默认模型（例如选了国内却写成 gpt-4o 主路径）。",
    ]
    n = 10
    if do_ft:
        rules.append(
            f"{n}. **模型微调（开启，不可省略）**：\n"
            "   - 明确基座模型（Qwen/DeepSeek/Llama 等，或 API 主路径 + 开源小模型旁路微调）。\n"
            "   - 范式：SFT / LoRA / QLoRA（首选）/ 可选 DPO；数据规范、规模、评测对比表。\n"
            "   - 交付：`finetune/`、`data/`、`evals/`（基座 vs 微调）。\n"
            "   - 微调计入「后训练」覆盖项，并仍须凑满 ≥8 个核心技术栈条目 + LLM 覆盖表齐全。"
        )
        n += 1
    else:
        rules.append(
            f"{n}. **模型微调：已关闭** —— 不要输出微调专章；后训练覆盖项写「使用已后训练权重 + 推理侧策略」；"
            "预训练覆盖项写「复用开源预训练 checkpoint」；改进点侧重 RAG选型改进/Multi-Agent/记忆/流式/部署等推理与工程侧优化。"
        )
        n += 1
    if do_gh and do_demo:
        rules.append(
            f"{n}. **部署交付（GitHub + 在线 Demo 均开启）**：\n"
            "   - A. GitHub 公开仓库（完整代码 + README）\n"
            "   - B. 在线 Demo 必须**国内网络默认可打开**（面试官无需科学上网）。\n"
            "   - **禁止平台**：Hugging Face Spaces、Streamlit Cloud、Vercel、Netlify、Railway、Render、Heroku 等需科学上网或不稳定的海外托管。\n"
            "   - **优先平台**：阿里云/腾讯云/华为云轻量或容器、Sealos、魔搭创空间、自建公网 IP+Nginx、（演示应急）cpolar/花生壳等国内可访问方案。\n"
            "   - 提示词：可运行 MVP → 推 GitHub → 国内可访问 Demo；README 顶部放 Demo 链接。"
        )
        n += 1
    elif do_gh:
        rules.append(
            f"{n}. **GitHub 公开仓（开启；在线 Demo 关闭）**：\n"
            "   - 必须写清仓库结构、README、License、推送步骤。\n"
            "   - **不要**要求公网在线 Demo / 云托管上线。"
        )
        n += 1
    elif do_demo:
        rules.append(
            f"{n}. **在线 Demo（开启；GitHub 推仓可弱化）**：\n"
            "   - Demo 必须**国内网络默认可打开**。\n"
            "   - **禁止平台**：HF Spaces、Streamlit Cloud、Vercel、Netlify、Railway、Render、Heroku。\n"
            "   - **优先**：阿里云/腾讯云/华为云轻量、Sealos、魔搭创空间、自建公网。\n"
            "   - 代码可本地交付；不要强制「必须先推公开 GitHub」作为硬性验收（除非用户另行要求）。"
        )
        n += 1
    else:
        rules.append(
            f"{n}. **GitHub / 在线 Demo：均已关闭** —— 不要要求公网 Demo / 强制推仓；本地可运行 + README 即可。"
        )
        n += 1

    todo_extra = []
    if do_ft:
        todo_extra.append("微调")
    if do_gh:
        todo_extra.append("GitHub")
    if do_demo:
        todo_extra.append("在线 Demo")
    todo_extra_s = "、".join(todo_extra) if todo_extra else "无额外项"

    rules.append(
        f"{n}. 「AI 自动化搭建提示词」：**每个项目恰好 1 份**可复制 ```text 搭建指令（只用于 Coding Agent 落地）；"
        f"**不要**再写其它提示词（如面试话术提示词、作品集包装提示词、第二段续写提示词等）。"
        f"提示词必须写清**端到端项目流程**（按 LLM 全链路阶段）、**LLM 主流技术覆盖表**、**Agent-Skill 映射表**与 **创新点**（≥2 条，含相对传统方案的↑/↓%）。"
        f"TODO 须含：应用本体、**≥8 技术栈落地**、**LLM 全链路各阶段可运行切片**、**改进点对照评测**"
        + (f"、{todo_extra_s}" if todo_extra else "")
        + "；单任务通关；禁止虚构业绩。"
        + ("非同质化差异化业务闭环。" if do_diff else "")
    )
    n += 1
    rules.append(
        f"{n}. 输出中文 Markdown；全文恰好 **{proj_n}** 个 AI 搭建提示词 fenced code block（每个项目 1 个，不要多不要少）。"
    )
    n += 1
    if do_iv:
        rules.append(
            f"{n}. **面试相关问题（已开启，不可省略）**：\n"
            "   - 每个项目必须有独立专章「面试相关问题」或「面试追问」。\n"
            f"   - {'简短版：3～5 题' if is_brief else '详细版：6～8 题'}；每题含：问题 + 简要回答框架（要点 bullet，不必长文）。\n"
            "   - 问题须紧扣本项目技术栈、自研改进、指标与流程，避免空泛八股。\n"
            "   - **禁止**把面试题写进「交给其他 AI 的搭建提示词」代码块。"
        )
        n += 1
    else:
        rules.append(
            f"{n}. **面试相关问题：已关闭** —— 禁止输出面试追问 / 模拟面试 Q&A。"
        )
        n += 1
    if do_mm:
        rules.append(
            f"{n}. **多模态（已开启，不可省略）**：\n"
            "   - 每个项目主路径必须包含至少一种非纯文本模态：图像 / 文档版面截图 / 音视频帧 等（按 JD 与业务选型）。\n"
            "   - 写清：输入形态、多模态模型或 API（视觉/语音等）、与 RAG/Agent/工具调用的衔接、评测样例（含多模态案例数据）。\n"
            "   - 技术栈与 LLM 覆盖表中须体现多模态条目；相对「纯文本基线」至少 1 条可量化改进（↑/↓%）。\n"
            + ("   - 多模态要服务具体业务闭环，尽量避免同质化的通用问答壳。" if do_diff else "   - 多模态要服务具体业务闭环即可。")
        )
        n += 1
    else:
        rules.append(
            f"{n}. **多模态：已关闭** —— 默认纯文本/结构化文本链路即可；不要硬塞视觉/语音模型或无关多模态栈。"
        )
        n += 1
    if do_iv:
        _ban_extra = (
            "   - **禁止写**：作品集包装专章/附录、额外叙事附录；面试题只写在正文专章。\n"
            "   - 搭建提示词里不要写「面试追问」；只保留搭建执行内容。"
        )
    else:
        _ban_extra = (
            "   - **禁止写**：面试追问、模拟面试 Q&A、作品集包装专章/附录、额外叙事附录。\n"
            "   - 提示词里也不要写「面试追问」；只保留搭建执行内容。"
        )
    rules.append(
        f"{n}. **字数与禁写硬性要求**：\n"
        f"   - 每个项目整体 ≤ **15000 字**。\n"
        + _ban_extra
    )
    n += 1
    rules.append(
        f"{n}. **每个项目的 AI 搭建提示词结构（```text 块内）**：只写搭建相关，按下列标题写满：\n"
        "   1) 【角色与目标】\n"
        "   2) 【项目做什么】一句话定位 + 差异化\n"
        "   3) 【端到端流程】编号步骤须体现 LLM 全链路：\n"
        "      数据/知识 →（预训练权重/CPT）→ 后训练或选用后训练权重 → RAG（按业务选型，勿一律 GraphRAG） → Multi-Agent → Function Calling/MCP → 记忆 → LLM推理(含RoPE/长上下文) → 流式输出 → 部署/交付 → 评测"
        + ("（含微调目录）" if do_ft else "（微调关闭则后训练写选用现成权重）")
        + (" → GitHub" if do_gh else "")
        + (" → 国内 Demo" if do_demo else "")
        + "；每步写清产物\n"
        "   4) 【LLM 主流技术覆盖表】逐项覆盖：LLM、RoPE、RAG（按业务选型）、Multi-Agent、Function Calling、MCP、记忆系统、流式输出、模型部署、后训练"
        + ("、预训练/CPT" if do_pt else "（预训练关闭则不写 CPT）")
        + "（阶段+做法+相对传统改进）\n"
        "   5) 【Agent-Skill 映射】每个 Agent 角色绑定对应 skill，写清输入/输出/工具/MCP/失败回退/验收指标\n"
        "   6) 【创新点】≥2 条：传统基线 → 做法 → **↑/↓%** → 测法\n"
        "   7) 【指标对比验证】基线 vs 改进（含百分比）\n"
        "   8) 【评测数据】案例数据 + 真实数据导入\n"
        "   9) 【技术栈清单】≥8 项，并标注对应流程阶段\n"
        "   10) 【强制执行清单 / TODO】单任务通关\n"
        "   11) 【验收清单】可勾选\n"
        "   **不要**在提示词里写【面试 STAR 摘要】或【面试追问】。"
    )

    overview_cols = "| 序号 | 定位 | 项目名 | 形态 | 差异化卖点 | 核心栈数 | LLM全链路 | 改进亮点（含↑/↓%） |"
    overview_sep = "|:----:|:----:|--------|:----:|------------|:--------:|:---------:|----------------------|"
    overview_extra_hdr = ""
    overview_extra_sep = ""
    overview_extra_cell = ""
    if do_ft:
        overview_extra_hdr += " 后训练/微调 |"
        overview_extra_sep += "--------------|"
        overview_extra_cell += " LoRA/QLoRA/... |"
    if do_gh:
        overview_extra_hdr += " GitHub |"
        overview_extra_sep += "---------|"
        overview_extra_cell += " 公开仓 |"
    if do_demo:
        overview_extra_hdr += " Demo 平台（国内可访问） |"
        overview_extra_sep += "-------------------------|"
        overview_extra_cell += " 阿里云轻量/Sealos/... |"
    overview_cols += overview_extra_hdr + " 建议工期 |"
    overview_sep += overview_extra_sep + "----------|"
    overview_rows = "\n".join(
        f"| {p['idx']} | {p['tier']} | ... | {form_label} | ... | ≥8 | 已覆盖 | ... |{overview_extra_cell} {timeline_label} |"
        for p in plan
    )

    stack_rows = [
        "| 前端 | ... | ... | ... |",
        "| 后端 | ... | ... | ... |",
        f"| 推理 API | {provider} / {model}（{region_label}） | 按用户选型 | 擅自换路线 |",
    ]
    if do_ft:
        stack_rows.append("| 微调 | 基座+LoRA/QLoRA+... | ... | ... |")
        stack_rows.append("| 评测 | ... | ... | ... |")
    if do_gh:
        stack_rows.append("| 仓库 | GitHub 公开 | 面试官可看 | 私有仓 |")
    if do_demo:
        stack_rows.append("| 在线 Demo | 阿里云/腾讯云/Sealos/魔搭/自建公网 | 国内可访问 | HF/Vercel 等需翻墙 |")
    if not do_gh and not do_demo:
        stack_rows.append("| 本地交付 | README + 一键启动 | 可演示 | 强制上线 |")

    sec_ft = ""
    if do_ft:
        sec_ft = """
## 1.4 后训练 / 模型微调方案（面试深度；对应「后训练」覆盖项）
- 预训练权重来源（复用开源 checkpoint；可选小规模领域 CPT，禁止从零预训练）：
- 基座模型与选型理由：
- 数据：来源/字段 schema/规模/清洗规则：
- 后训练方法：SFT 或 LoRA/QLoRA（写清秩、学习率量级、epoch 建议区间）；可选 DPO：
- 训练环境：必须贴合用户 GPU/CUDA/Python；不够则降级：
- 评测：基座/传统基线 vs 后训练后对比指标（≥2 个，含↑/↓%）：
- 与线上推理 / Multi-Agent / Function Calling 的衔接：
"""
    else:
        sec_ft = """
## 1.4 后训练 / 模型微调
（本次关闭）明确写：不进行微调训练；「后训练」覆盖项 = 使用已后训练权重 + 推理侧对齐/偏好策略；推理主路径用所选 API/本地服务。
"""

    sec_dp_parts = []
    if do_gh or do_demo:
        sec_dp_parts.append("## 1.5 部署交付（按开关）")
        if do_gh:
            sec_dp_parts.append(
                "### A. GitHub\n"
                "- 仓库结构要点、README 必含章节、License"
                + ("、Demo 链接位置" if do_demo else "")
            )
        else:
            sec_dp_parts.append("### A. GitHub\n（本次关闭）不要强制推公开仓。")
        if do_demo:
            sec_dp_parts.append(
                "### B. 在线 Demo（国内默认可访问；禁止需科学上网的平台）\n"
                "- 平台择一：阿里云/腾讯云/华为云轻量或容器、Sealos、魔搭创空间、自建公网 IP+Nginx；（应急）cpolar/花生壳\n"
                "- **禁止**：Hugging Face Spaces、Streamlit Cloud、Vercel、Netlify、Railway、Render、Heroku\n"
                "- 部署步骤、环境变量清单、公网验收标准（国内网络可打开主路径）"
            )
        else:
            sec_dp_parts.append("### B. 在线 Demo\n（本次关闭）不要要求公网上线。")
        sec_dp = "\n" + "\n".join(sec_dp_parts) + "\n"
    else:
        sec_dp = """
## 1.5 部署
（本次关闭）只需本地可运行说明；不要写强制上线/推仓步骤。
"""

    metric_extra = [
        "| 性能或轻量化 | 基线… → 改进…（须写 ↓xx% 或 ↑xx%） | 同脚本对照 |",
        "| 痛点相关指标 | 基线… → 改进…（须写 ↑/↓ 百分比） | 案例集 + 真实数据各跑一轮 |",
    ]
    if do_ft:
        metric_extra.append("| 微调增益 | 基座… → 微调后…（↑/↓%） | 同测评集对比表 |")
    if do_demo:
        metric_extra.append("| Demo 可用性 | 公网可访问 | 面试官浏览器实测 |")
    if do_gh and not do_demo:
        metric_extra.append("| 仓库可读性 | 公开可 clone + README 可复现 | 面试官实测 |")
    metric_extra_s = "\n".join(metric_extra)

    star_hooks = [
        "≥8 技术栈",
        "LLM全链路（Multi-Agent/RAG按业务选型/MCP/记忆/流式等）",
        "Agent-Skill 映射（每个 Agent 有对应 skill）",
        "差异化业务痛点",
        "自研改进≥2（各含目标↑/↓%）",
        "较新栈写清飞书手册中 RAG/Agent/MCP/记忆/部署等具体新点（按业务选型，勿一律硬塞）",
        "案例+真实数据双通道评测",
    ]
    if do_ft:
        star_hooks.append("后训练/微调取舍")
    if do_gh:
        star_hooks.append("GitHub 交付取舍")
    if do_demo:
        star_hooks.append("国内可访问 Demo 取舍")
    star_hook_s = " + ".join(star_hooks)

    task_must = [
        "①可运行应用",
        "②≥8 核心技术栈落地",
        "③LLM 主流技术全链路覆盖并按流程串联",
        "④Agent-Skill 映射表",
        "⑤相对传统方案的改进对照评测（输出↑/↓百分比）",
        "⑥案例数据一键评测 + 真实数据导入评测",
        "⑦差异化业务闭环（尽量避免同质化）" if do_diff else "⑦按 JD 落地的业务闭环（可含客服/问答等常规方向）",
    ]
    if do_ft:
        task_must.append("⑧后训练/模型微调流水线")
    if do_gh:
        task_must.append("⑨GitHub 公开仓")
    if do_demo:
        task_must.append("⑩国内可访问在线 Demo")
    task_must_s = " ".join(task_must)

    action_lines = [
        "1) docs：PRD.md / DESIGN.md / ARCHITECTURE.md / TODO.md / PIPELINE.md / LLM_STACK.md / AGENT_SKILLS.md / IMPROVEMENTS.md / METRICS.md"
        + (" / DEPLOY.md" if do_dp else ""),
        "2) 先写清 LLM 全链路流程再写代码：数据/知识 → "
        + ("预训练/CPT → " if do_pt else "（跳过 CPT，复用预训练权重）→ ")
        + "后训练或选用后训练权重 → "
        + ("多模态理解 → " if do_mm else "")
        + "RAG（按业务选型，勿一律 GraphRAG） → Multi-Agent → Function Calling/MCP → 记忆 → LLM推理(RoPE/长上下文) → 流式输出 → 部署 → 评测；PIPELINE.md + LLM_STACK.md 逐步落地",
        "3) 技术栈：写死 ≥8 个核心技术/方法（较新栈≥3：须写飞书手册中较新且贴业务的具体进展，禁止一律硬塞；自研改进≥2），并填满 LLM 主流技术覆盖表与 Agent-Skill 映射表"
        + ("（含预训练/CPT）" if do_pt else "（预训练关闭，不写 CPT）")
        + "；禁止只有 CRUD；禁止空名词堆砌"
        + ("；尽量避免与市面 Demo 同质化" if do_diff else "；可做客服/问答等常规方向"),
        "4) 自研改进/创新点：硬性 ≥2 条写入 IMPROVEMENTS.md（传统基线 / 做法 / 目标↑或↓百分比 / 测法），并在代码里真实现、可跑通对照实验",
        "5) 评测双通道：data/samples/ 内置案例一键跑；另提供真实数据上传/导入（UI 或 CLI），同一套 eval 输出对照表到 METRICS.md",
        "6) 业务：无公司名；用业务对象+流程+痛点驱动功能设计",
        "7) 应用实现：目录、核心接口与页面、三态 UI、SSE 流式；主路径按流程逐步可演示；页面上能切换「跑案例」与「导入真实数据」；每个 Agent 在代码里有对应 skill/handler",
        "8) 改进评测：传统/朴素基线 vs 改进版；每个指标输出 基线→改进→变化百分比；STAR-R 必须引用这些百分比",
    ]
    ai = 9
    if do_ft:
        action_lines.append(
            f"{ai}) 微调：finetune/ + data/ + 训练命令 + evals/；按用户 GPU/CUDA 降级"
        )
        ai += 1
        action_lines.append(f"{ai}) 评测：本地黄金样例 + 可导入真实集，metrics.md 含百分比")
        ai += 1
    if do_gh:
        action_lines.append(
            f"{ai}) GitHub：README"
            + ("（顶部 Demo URL 占位）" if do_demo else "")
            + "、.env.example、License、推送步骤；README 写清案例/真实数据两种测法"
        )
        ai += 1
    if do_demo:
        action_lines.append(
            f"{ai}) 在线 Demo：阿里云/腾讯云/华为云轻量、Sealos、魔搭创空间或自建公网（禁止 HF Spaces/Vercel/Railway/Streamlit Cloud 等需科学上网平台）；验收=国内网络可打开主路径；Demo 上可跑案例并支持上传真实数据"
        )
        ai += 1
    action_lines.append(f"{ai}) 每次只做 TODO 一条；通关后再下一条")
    action_s = "\n".join(action_lines)

    result_checks = [
        "- [ ] 本地一键启动成功",
        "- [ ] PIPELINE.md / LLM_STACK.md 与代码主路径一致（技术按流程串联）",
        "- [ ] AGENT_SKILLS.md 含 Agent-Skill 映射表，每个 Agent 均绑定 skill、输入输出、工具/MCP、失败回退与验收指标",
        "- [ ] 核心技术栈清单 ≥8；其中较新栈≥3（飞书手册具体新点+业务选型理由，非空名词）",
        "- [ ] 自研改进 ≥2 条（技术栈表类型=自研改进 + 改进对照表≥2 行，均含↑/↓%）",
        "- [ ] LLM 覆盖表齐全：LLM / RoPE / RAG（按业务选型） / Multi-Agent / Function Calling / MCP / 记忆 / 流式 / 部署 / 后训练"
        + (" / 预训练(CPT)" if do_pt else "（预训练开关关闭，不要求 CPT）")
        + (" / 多模态" if do_mm else ""),
        "- [ ] 创新点 ≥2 条已实现；METRICS/对照表含传统基线→改进→↑/↓百分比",
        "- [ ] 案例数据可一键评测；真实数据可通过上传/导入跑同一套评测",
        "- [ ] STAR-R 能用百分比说明「因某设计而提升/降低多少」",
        (
            "- [ ] 有差异化业务闭环，尽量避免与市面常见 Demo 同质化"
            if do_diff
            else "- [ ] 有可运行的业务闭环（差异化关闭：可含客服/问答等常规方向）"
        ),
        "- [ ] 业务描述无公司名",
    ]
    if do_mm:
        result_checks.append("- [ ] 多模态主路径可演示（非纯文本输入 + 多模态模型/API + 评测样例）")
    if do_ft:
        result_checks.append("- [ ] 微调目录与对比表齐全（或诚实降级说明）")
    if do_gh:
        result_checks.append("- [ ] GitHub 公开可读可 clone")
    if do_demo:
        result_checks.append("- [ ] 在线 Demo 国内网络可打开主路径（非海外需翻墙平台）")
    result_checks.append("- [ ] 指标表含基线/改进/变化百分比与测量方法；无编造数字")
    result_s = "\n".join(result_checks)

    # AI 搭建提示词内：端到端流程步骤（按开关编号，避免跳号）
    pipeline_steps = [
        "1. 数据/知识构建：来源、格式、样例路径（data/raw/、图谱/语料"
        + ("、多模态样例如图片/音视频" if do_mm else "")
        + "）→ 产物：…",
        "2. 预训练策略："
        + (
            "开启 CPT/继续预训练（小规模，按显存降级；禁止从零预训练）"
            if do_pt
            else "关闭 CPT：直接复用开源预训练 checkpoint（一句即可，不写训练专章）"
        )
        + " → 产物：…",
        "3. 后训练："
        + ("SFT/LoRA/QLoRA（可选 DPO）流水线 + 数据规范" if do_ft else "使用已后训练权重 + 推理侧对齐/偏好策略（微调关闭，禁止假装训练）")
        + " → 产物：…",
    ]
    if do_mm:
        pipeline_steps.append(
            "4. 多模态理解：图像/音视频等输入预处理 + 视觉/多模态模型或 API 调用 → 产物：…"
        )
        _mm_n = 5
    else:
        _mm_n = 4
    pipeline_steps.extend(
        [
            f"{_mm_n}. RAG 索引与检索（按业务选合适且效果更好的方案；手册方法仅参考，不强制 Hybrid/RAPTOR/CRAG/GraphRAG）→ 产物：…",
            f"{_mm_n + 1}. Multi-Agent 编排：角色分工、握手协议、失败回退 → 产物：…",
            f"{_mm_n + 2}. Function Calling + MCP：工具 schema、MCP Server、参数校验与错误回灌 → 产物：…",
            f"{_mm_n + 3}. 记忆系统：短期工作记忆 + 长期记忆读写策略 → 产物：…",
            f"{_mm_n + 4}. LLM 推理：按 {provider}/{model}（{region_label}）接入；含 RoPE/长上下文策略、超时重试"
            + ("；与多模态结果融合" if do_mm else "")
            + " → 产物：…",
            f"{_mm_n + 5}. 流式输出：SSE/流式 token 推到前端；三态 UI；支持「跑内置案例」与「导入真实数据」→ 产物：…",
            f"{_mm_n + 6}. 模型部署：本地推理服务或按开关的公开仓/国内 Demo → 产物：…",
            f"{_mm_n + 7}. 评测对照：传统基线 vs 改进；输出基线→改进→↑/↓% → 产物：evals/、METRICS.md；案例集与真实数据各跑一轮",
            f"{_mm_n + 8}. 文档：PIPELINE.md、LLM_STACK.md、IMPROVEMENTS.md、METRICS.md、ARCHITECTURE.md、TODO.md",
        ]
    )
    _ps = _mm_n + 9
    if do_ft:
        pipeline_steps.append(f"{_ps}. 后训练/微调目录落地：finetune/ + data/ + evals/ → 产物：…")
        _ps += 1
    if do_gh:
        pipeline_steps.append(
            f"{_ps}. GitHub 公开仓：README"
            + ("（顶部 Demo URL）" if do_demo else "")
            + "、License、推送 → 产物：公开仓库"
        )
        _ps += 1
    if do_demo:
        pipeline_steps.append(f"{_ps}. 国内可访问 Demo 上线与验收 → 产物：公网可打开主路径")
    pipeline_s = "\n".join(pipeline_steps)

    brief_pipeline = (
        "1. 数据/知识构建"
        + ("（含多模态样例）" if do_mm else "")
        + "：…\n"
        + (
            "2. 预训练/CPT + 后训练策略：…\n"
            if do_pt
            else "2. 后训练策略（预训练关闭：复用开源权重，不写 CPT）：…\n"
        )
        + ("3. 多模态理解链路：…\n" if do_mm else "")
        + f"{'4' if do_mm else '3'}. RAG（按业务选型） + Multi-Agent + Function Calling/MCP + 记忆：…\n"
        f"{'5' if do_mm else '4'}. LLM 推理（{provider}/{model}，含 RoPE/长上下文）+ 流式输出：…\n"
        f"{'6' if do_mm else '5'}. 模型部署与界面：…\n"
        f"{'7' if do_mm else '6'}. 评测对照（传统基线 vs 改进）：…\n"
        f"{'8' if do_mm else '7'}. 交付："
        + "、".join(
            ["本地可跑"]
            + (["多模态"] if do_mm else [])
            + (["预训练/CPT"] if do_pt else [])
            + (["后训练/微调"] if do_ft else [])
            + (["GitHub"] if do_gh else [])
            + (["国内 Demo"] if do_demo else [])
        )
    )

    omit_note = []
    if not do_ft:
        omit_note.append("不要写微调专章")
    if not do_pt:
        omit_note.append("不要写预训练/CPT 专章（一句复用开源权重即可）")
    if not do_gh:
        omit_note.append("不要强制 GitHub 推仓专章")
    if not do_demo:
        omit_note.append("不要写在线 Demo/公网上线专章")
    if omit_note:
        omit_s = "；".join(omit_note)
    else:
        omit_s = "微调、预训练、GitHub、在线 Demo 均需按开启项写全"

    pipeline_extra_steps = []
    step_n = 13
    if do_ft:
        pipeline_extra_steps.append(f"{step_n}. 后训练/微调流水线：finetune/ + data/ + evals/ → 产物：…")
        step_n += 1
    if do_gh:
        pipeline_extra_steps.append(
            f"{step_n}. GitHub 公开仓：README"
            + ("（顶部 Demo URL）" if do_demo else "")
            + "、License、推送 → 产物：公开仓库"
        )
        step_n += 1
    if do_demo:
        pipeline_extra_steps.append(f"{step_n}. 国内可访问 Demo 上线与验收 → 产物：公网可打开的主路径")
        step_n += 1
    pipeline_extra_s = ("\n" + "\n".join(pipeline_extra_steps)) if pipeline_extra_steps else ""

    brief_deliver = "本地可跑"
    if do_ft:
        brief_deliver += " / 微调"
    if do_gh:
        brief_deliver += " / GitHub"
    if do_demo:
        brief_deliver += " / 国内 Demo"

    overview_tail_ft = " ..." if do_ft else ""
    overview_tail_gh = " ..." if do_gh else ""
    overview_tail_demo = " ..." if do_demo else ""
    demo_access_note = "Demo 须国内可访问" if do_demo else "本次不要求在线 Demo"
    diff_note = "差异化选题（尽量避免同质化）" if do_diff else "可含客服/问答等常规方向（差异化关闭）"
    max_tok = 4096 if is_brief else 8192
    # 按项目数粗调上限
    max_tok = min(12288, max_tok + max(0, proj_n - 1) * 2048)

    first = plan[0]
    more_projects_brief = ""
    more_projects_full = ""
    for p in plan[1:]:
        more_projects_brief += (
            f"\n# 🚀 {p['title']}<项目名>\n"
            f"**完整写出**与项目1同级小节（含 STAR + **恰好 1 份**搭建提示词 ```text）；"
            f"本项目整体 ≤15000 字；禁止同上/略；禁止作品集专章"
            + ("；须含面试相关问题" if do_iv else "；禁止面试追问")
            + "。\n"
        )
        more_projects_full += (
            f"\n# 🚀 {p['title']}<项目名>\n"
            f"**必须完整展开**与项目1同级小节（{p['idx']}.1～{p['idx']}.8），"
            f"含完整 STAR，以及 **恰好 1 份**「交给其他 AI 的自动化搭建提示词」```text；"
            f"**禁止**结构同上/略；**禁止**作品集附录"
            + ("；**必须**含面试相关问题专章" if do_iv else "；**禁止**面试追问")
            + f"；本项目整体 ≤15000 字。\n"
            f"（≥8 技术栈 + LLM 全链路覆盖；{diff_note}；无公司名；遵守范围开关与 API 选型）\n"
        )

    _llm_chain = (
        "LLM/RoPE/RAG按业务选型/Multi-Agent/Function Calling/MCP/记忆/流式/部署/后训练"
        + ("/预训练" if do_pt else "")
    )
    _iv_head = (
        "面试相关问题写在正文专章（题目+回答框架），**不要**写进搭建提示词"
        if do_iv
        else "**不要**面试追问"
    )
    if do_iv:
        sec_iv_brief = """
## 面试相关问题（3～5 题；紧扣本项目技术/改进/指标）
1. **Q：** …  
   **回答框架：** …
2. **Q：** …  
   **回答框架：** …
3. **Q：** …  
   **回答框架：** …
（可再补 1～2 题）
"""
        sec_iv_full = """
## 1.9 面试相关问题（6～8 题；紧扣本项目技术栈、自研改进、指标与流程）
1. **Q：** …  
   **回答框架：** …
2. **Q：** …  
   **回答框架：** …
3. **Q：** …  
   **回答框架：** …
4. **Q：** …  
   **回答框架：** …
5. **Q：** …  
   **回答框架：** …
6. **Q：** …  
   **回答框架：** …
（可再补 1～2 题；不要空泛八股）
"""
    else:
        sec_iv_brief = ""
        sec_iv_full = ""

    common_head = f"""你是「Tech Lead + 项目方案架构师」。根据目标岗位 JD，推荐恰好 **{proj_n}** 个可在 **{timeline_label}** 内用 Coding Agent 落地的**{scale_label}·{form_label}**（质量：{tier_names}）。
每个项目必须：{must_line}。
篇幅：**{detail_label}**；每项目硬上限 **≤15000 字**；工期：**{timeline_label}**；背景：**{scale_label}**；形态：**{form_label}**。
每个项目写 **恰好 1 份**「交给其他 AI 的自动化搭建提示词」（```text，只用于搭建）；**不要**其它提示词、**不要**作品集专章；{_iv_head}。
搭建提示词主体 = 按真实工程流程串联的 LLM 全链路（{_llm_chain}）+ 相对传统方案的创新点（含↑/↓%）+ 案例/真实数据双通道 + 可执行 TODO。

## 目标岗位 JD
{req.jd}{bg}{pref}{env}{llm_scope}{scope}{plan_scope}{length_scope}

---

## 硬性要求（必须遵守）
{chr(10).join(rules)}

{_FEISHU_TECH_ROUTE}

---
"""

    if is_brief:
        brief_extra = []
        if do_ft:
            brief_extra.append("- 微调：基座 + 方法 + 数据规模 + 评测各 1 行")
        if do_pt:
            brief_extra.append("- 预训练/CPT：数据规模 + 算力约束 + 验收（3～5 行）")
        if do_gh:
            brief_extra.append("- GitHub：仓库要点 + README 必含项（3～5 行）")
        if do_demo:
            brief_extra.append("- Demo：国内可访问平台 + 验收标准（3～5 行）")
        brief_extra_s = "\n".join(brief_extra) if brief_extra else "- （无额外部署/微调/预训练专章）"
        _cov_brief = (
            "LLM / RoPE / RAG（按业务选型） / Multi-Agent / Function Calling / MCP / 记忆 / 流式 / 部署 / 后训练"
            + (" / 预训练" if do_pt else "（预训练关）")
            + (" / 多模态" if do_mm else "（多模态关）")
        )
        _pt_flag = "开" if do_pt else "关"
        _mm_flag = "开" if do_mm else "关"
        _cov_list_brief = (
            "LLM｜RoPE｜RAG（按业务选型）｜Multi-Agent｜Function Calling｜MCP｜记忆系统｜流式输出｜模型部署｜后训练"
            + ("｜预训练/CPT" if do_pt else "（预训练关，不写 CPT）")
            + ("｜多模态" if do_mm else "（多模态关）")
        )

        prompt = common_head + f"""
# 📦 推荐项目一览（简短）

{overview_cols}
{overview_sep}
{overview_rows}

---

# 🚀 {first['title']}<项目名>

## 定位 + 差异化（2～3 句）

## 痛点清单（恰好 3 条，无公司名）
1. ...
2. ...
3. ...

## 核心技术栈（≥8）+ 自研改进（自研改进≥2；较新栈≥3 须写具体新点）
| # | 技术/方法 | 类型 | 流程阶段 | 一句话用途/相对传统改进 |
|:-:|-----------|------|----------|-------------------------|
| 1-8 | ... | 基础工程 / 较新栈（飞书手册新点+业务理由） / 自研改进 | 数据/检索/编排/推理/部署… | ... |

### LLM 主流技术覆盖（简表；预训练按开关）
| 技术 | 流程阶段 | 做法一句话 | 相对传统改进（含%或开关等价落地） |
|------|----------|------------|-----------------------------------|
| {_cov_brief} | … | … | … |

### Agent-Skill 映射（每个 Agent 必须配对应 skill）
| Agent 角色 | 对应 skill | 触发时机 | 输入 | 输出 | 工具/MCP/函数 | 失败回退 | 验收指标 |
|------------|------------|----------|------|------|---------------|----------|----------|
| 规划 Agent | task_planning_skill | … | … | … | … | … | … |
| 检索/图谱 Agent | graph_retrieval_skill | … | … | … | … | … | … |
| 工具执行/评测 Agent | mcp_tool_calling_skill / eval_metric_skill | … | … | … | … | … | … |

改进对照（2 行）：传统基线值 → 做法 → 改进值 → ↑/↓百分比 → 测法（案例集 + 真实数据）

{brief_extra_s}

## 面试可讲的完整 STAR（简短版也禁止略写）
### S Situation（≥40 字）
### T Task（≥30 字）
### A Action（≥3 条；含 {star_hook_s}；API={provider}/{model}）
### R Result（≥2 条；必须写「因某设计 → 某指标 ↑xx% 或 ↓xx%」；案例+真实数据均可复测）
{sec_iv_brief}
## 交给其他 AI 的精简搭建提示词
```text
【角色与目标】用 {provider}/{model}（{region_label}）落地本项目。范围：微调={'开' if do_ft else '关'}；预训练={_pt_flag}；多模态={_mm_flag}；面试题={'开' if do_iv else '关'}；GitHub={'开' if do_gh else '关'}；Demo={'开' if do_demo else '关'}。
必须：{task_must_s}；≥8 技术栈；LLM 全链路覆盖；{diff_note}；无公司名；{demo_access_note}；{omit_s}。

【项目做什么】（1～2 句定位 + {'差异化卖点' if do_diff else '定位说明（差异化卖点可选）'}；形态必须是「{form_label}」；禁止只写空泛「做一个 AI 系统」）

【端到端流程】（必须按 LLM 工程流程编号，不能省略）
{brief_pipeline}
每步写清产物目录或文件名。

【LLM 主流技术覆盖表】（必写：阶段 + 做法 + 相对传统改进；预训练按开关）
{_cov_list_brief}

【Agent-Skill 映射】（必写；设计 Agent 时要搭配对应 skill）
- 规划 Agent → task_planning_skill：输入=…；输出=…；工具/MCP=…；失败回退=…；验收=…
- 检索/图谱 Agent → graph_retrieval_skill：输入=…；输出=…；工具/MCP=…；失败回退=…；验收=…
- 工具执行 Agent → mcp_tool_calling_skill：输入=…；输出=…；工具/MCP=…；失败回退=…；验收=…
- 评测 Agent → eval_metric_skill：输入=…；输出=METRICS.md/对照表；工具/MCP=…；失败回退=…；验收=…

【创新点】（≥2 条，禁止空话；相对传统基线；收益必须写百分比）
- 创新1：传统基线=… → 做法=… → 目标↑或↓xx% → 测法=…
- 创新2：传统基线=… → 做法=… → 目标↑或↓xx% → 测法=…

【指标对比验证】
- 对照表：指标 | 传统基线值 | 改进值 | 变化百分比 | 测法
- STAR-R 话术：因为做了 A 中的…设计，所以…指标 ↑/↓ xx%

【评测数据双通道】
- 案例：data/samples/ 一键跑通对照
- 真实：UI 上传或 CLI 指定路径导入用户数据，同一套 eval 出百分比表

【技术栈】列出 ≥8 项，标注流程阶段，并标出哪几项对应上述创新点。

【执行】先写 PIPELINE.md + LLM_STACK.md + IMPROVEMENTS.md + METRICS.md，再按 TODO 单任务通关；验收含本地可跑 + 全链路流程跑通 + 百分比对照表 + 真实数据可导入。
（本块是唯一搭建提示词；不要追加面试追问或其它提示词）
```

---
{more_projects_brief}
---

# ✅ 选题与交付建议（各 1～2 句；不要作品集专章）
- 只做一个时选哪个
- {'如何避免同质化' if do_diff else '常规方向如何仍讲出工程亮点'}
- 简历一行示例（无公司名）
"""
    else:
        _pt_row = (
            "| 预训练 | 权重/CPT | 小规模领域 CPT（禁止从零预训练） | 无领域适应 / 声称从零预训练 | … |"
            if do_pt
            else "| 预训练 | （本次关闭） | 直接复用开源预训练 checkpoint，不写 CPT | — | 开关关闭，一句带过即可 |"
        )
        _pt_ai_item = (
            "11. 预训练/CPT（已开启：小规模 CPT；禁止从零预训练）"
            if do_pt
            else "11. 预训练（已关闭：复用开源 checkpoint，不写 CPT 专章）"
        )
        _mm_row = (
            "| 多模态 | 视觉/音视频理解 | 非纯文本输入 + 多模态模型/API + 与 Agent/RAG 衔接 | 纯文本基线 | … |"
            if do_mm
            else "| 多模态 | （本次关闭） | 纯文本/结构化文本链路即可 | — | 开关关闭，勿硬塞 |"
        )
        _mm_ai_item = (
            "12. 多模态（已开启：图像/音视频等输入 + 多模态模型/API + 评测样例）"
            if do_mm
            else "12. 多模态（已关闭：勿硬塞视觉/语音栈）"
        )
        _ft_ai = (
            "SFT/LoRA/QLoRA/可选DPO"
            if do_ft
            else "使用已后训练权重+推理侧策略；微调关闭禁止假装训练"
        )
        prompt = common_head + f"""
# 🎯 JD 技术栈锚定
（必备栈 / 加分栈 / 落地优先级）

# 📦 推荐项目一览

{overview_cols}
{overview_sep}
{overview_rows}

---

# 🚀 {first['title']}<项目名>
（本项目整体 ≤15000 字；禁止作品集附录；面试题开关={'开' if do_iv else '关'}）

## 1.1 一句话定位{' + 差异化卖点' if do_diff else '（差异化卖点可选）'}
- 定位：...
- {'相对常见 Demo（尤其客服/问答助手）的差异：...' if do_diff else '（可选）相对常见 Demo 的差异 / 或直接说明这是常规方向：...'}

## 1.2 业务背景（实际业务流程 + 痛点驱动；禁止任何公司名）
- 业务对象与流程环节：（不要出现公司/品牌名）
- 从 JD 抽象的能力需求（只写技能类型）：
- **痛点清单（≥3，要具体）**：
  1. ...
  2. ...
  3. ...
- 功能如何映射痛点：
- 约束（成本/时限/脏数据/并发/可观测性等，≥2）：

## 1.3 技术栈分层（工程落地）
| 层级 | 选型 | 选型理由（贴 JD） | 明确不做 |
|------|------|-------------------|----------|
{chr(10).join(stack_rows)}

## 1.3b 核心技术栈（必须 ≥8 项）+ 自研改进（≥2）+ 较新栈（≥3）
用表格列出至少 8 条。类型标注：`基础工程` / `较新栈` / `自研改进`。必须标注「流程阶段」。
- **较新栈**：须写飞书手册方向的**具体新进展**（RAG/Agent/MCP/记忆/推理部署等均可；**方法名按业务自选**，手册条目只是参考池），每条带**业务选型理由与效果预期**，禁止空名词与固定套路硬塞。
- **自研改进**：至少 2 条，且与下方改进对照表对应。

| # | 技术/方法 | 类型 | 流程阶段 | 用途 | 相对传统方案的改进 / 「新在哪里」 |
|:-:|-----------|------|----------|------|------------------------------|
| 1 | ... | ... | ... | ... | ... |
| 2 | ... | ... | ... | ... | ... |
| 3 | ... | ... | ... | ... | ... |
| 4 | ... | ... | ... | ... | ... |
| 5 | ... | ... | ... | ... | ... |
| 6 | ... | ... | ... | ... | ... |
| 7 | ... | ... | ... | ... | ... |
| 8 | ... | ... | ... | ... | ... |

## 1.3c LLM 主流技术覆盖表（按真实工程流程；预训练按开关；禁止只堆名词）
| 技术点 | 流程阶段 | 本项目具体做什么 | 传统/朴素做法 | 你的改进与收益（↑/↓% 或开关等价落地） |
|--------|----------|------------------|---------------|----------------------------------------|
| LLM | 推理主路径 | … | 无统一客户端/随意换模型 | … |
| RoPE / 长上下文 | 上下文策略 | … | 超长直接硬截断 | … |
| RAG（写清本业务选型与理由；手册方法仅参考，不强制某几种） | 知识检索 | … | 仅单路向量 TopK | … |
| Multi-Agent | 编排协作 | … | 单 Agent 一把梭 | … |
| Function Calling | 工具调用 | … | 模型空想直接答 | … |
| MCP | 工具协议面 | … | 临时脚本/无协议工具 | … |
| 记忆系统 | 状态保持 | … | 无状态每轮重传全文 | … |
| 流式输出 | 交互体验 | … | 整段阻塞返回 | … |
| 模型部署 | 可演示交付 | … | 仅本地 notebook | … |
| 后训练 | 对齐/微调 | {'SFT/LoRA/QLoRA（可选 DPO）' if do_ft else '使用已后训练权重 + 推理侧策略（微调关闭）'} | {'无后训练直接硬用基座' if do_ft else '假装做了训练（禁止）'} | … |
{_pt_row}
{_mm_row}

## 1.3d Agent-Skill 映射表（设计 Agent 时必须搭配对应 skill）
| Agent 角色 | 对应 skill | 触发时机 | 输入 | 输出 | 调用工具/MCP/函数 | 失败回退 | 验收指标 |
|------------|------------|----------|------|------|-------------------|----------|----------|
| 规划 Agent | task_planning_skill | … | JD/项目目标/约束 | 可执行计划/TODO | … | 计划不完整时回到需求澄清 | … |
| 检索 Agent | retrieval_skill（按实际选型命名） | … | 查询/文档 | 证据片段 | 按业务选型的检索/MCP 工具 | 检索为空时降级或补检索 | … |
| 工具执行 Agent | mcp_tool_calling_skill | … | schema 化任务参数 | 工具结果/错误回灌 | Function Calling/MCP Server | 参数缺失时补参重试 | … |
| 记忆 Agent | memory_update_skill | … | 会话/任务状态 | 短期/长期记忆记录 | 向量库/结构化存储 | 写入失败时保留工作记忆 | … |
| 评测 Agent | eval_metric_skill | … | 案例/真实数据/输出结果 | METRICS.md/对照表 | eval 脚本/指标函数 | 评测失败时输出失败样例 | … |

### 改进对照表（突出相对传统方案的性能/质量/成本，至少 2 行自研改进；收益必须含百分比）
| 改进点 | 传统/朴素基线（数值） | 你的做法 | 改进后（数值） | 变化百分比（↑/↓） | 测量方法 |
|--------|----------------------|----------|----------------|-------------------|----------|
| ... | ... | ... | ... | ↓xx% 或 ↑xx% | 案例集脚本 / 真实数据导入后同脚本 |
| ... | ... | ... | ... | … | … |

### 评测数据双通道
- **案例数据**：`data/samples/` 内置样例 + 一键命令/按钮；面试可稳定复现对照表。
- **真实数据**：页面上传或 CLI `--input` 指定路径；写清支持格式（CSV/JSON/…）、字段 schema、隐私（本地处理）。
- 两种数据源必须走**同一套**传统基线 vs 改进评测，输出到 `METRICS.md`。

{sec_ft}{sec_dp}
## 1.6 目标指标（Result）——用对比证明「更好」
| 指标 | 基线值 | 改进后 | 变化百分比 | 测量方法（案例 / 真实数据） |
|------|--------|--------|------------|------------------------------|
| ... | ... | ... | ↑xx% 或 ↓xx% | ... |
{metric_extra_s}

## 1.7 面试可讲的完整 STAR（禁止略写；四段都必须写满实质内容）
### S Situation（≥80 字）
（业务场景、痛点、约束；无公司名）

### T Task（≥60 字）
（你的目标、成功标准、范围边界；成功标准含「相对基线达到哪些 ↑/↓%」）

### A Action（≥5 条或 ≥150 字；必须覆盖：{star_hook_s}；并写明推理 API={provider}/{model}）
1. ...
2. ...
3. ...
4. ...
5. ...

### R Result（硬性：用指标对比验证设计；禁止空话）
每条写成：「因为 A 中做了… → 某指标从…到…（↑xx% 或 ↓xx%）；测法=案例集 / 真实数据导入」。
- ...
- ...
- （至少 2 条含明确百分比；说明案例与真实数据均可复测）
{sec_iv_full}
## 1.8 交给其他 AI 的自动化搭建提示词（直接复制）
```text
（必须写满；{omit_s}；≥8 核心技术栈；LLM 主流技术全链路覆盖；{diff_note}；业务无公司名；{demo_access_note}；推理API={provider}/{model}；面试题={'开' if do_iv else '关'}；多模态={'开' if do_mm else '关'}）
（硬性：提示词主体 = 按工程流程串联的端到端流程 + LLM 覆盖表 + 相对传统方案的创新点（含↑/↓%）+ 案例/真实数据双通道评测 + 可执行清单）

【角色与目标】
你是 Coding Agent / Tech Lead。按 SDD 单任务通关，交付可演示、可用指标百分比证明「比传统/朴素基线更好」的面试级作品；技术栈体现大模型项目完整性。

【项目做什么】
- 交付形态（硬性）：{form_label}
- 一句话定位：…
- {'差异化卖点（相对常见 Demo，尤其客服/问答）：…' if do_diff else '差异化卖点（可选；可写「常规方向」）：…'}
- 用户/业务对象与主场景（无公司名）：…

【端到端流程】（必须按 LLM 工程流程写具体，禁止「实现功能即可」；每步给产物）
{pipeline_s}

【LLM 主流技术覆盖表】（每项写：流程阶段 / 具体实现 / 相对传统改进；预训练按开关）
1. LLM
2. RoPE / 长上下文
3. RAG（按业务选型）
4. Multi-Agent
5. Function Calling
6. MCP
7. 记忆系统
8. 流式输出
9. 模型部署
10. 后训练（{_ft_ai}）
{_pt_ai_item}
{_mm_ai_item}

【Agent-Skill 映射】（必写；设计 Agent 时要搭配对应 skill）
请创建 `docs/AGENT_SKILLS.md`，并在代码中落地对应 handler/tool：
| Agent 角色 | 对应 skill | 触发时机 | 输入 | 输出 | 调用工具/MCP/函数 | 失败回退 | 验收指标 |
|------------|------------|----------|------|------|-------------------|----------|----------|
| 规划 Agent | task_planning_skill | … | … | … | … | … | … |
| 检索 Agent | hybrid_retrieval_skill（按选型） | … | … | … | 检索/MCP | … | … |
| 工具执行 Agent | mcp_tool_calling_skill | … | … | … | Function Calling/MCP Server | … | … |
| 记忆 Agent | memory_update_skill | … | … | … | 向量库/结构化记忆 | … | … |
| 评测 Agent | eval_metric_skill | … | … | METRICS.md/对照表 | eval 脚本 | … | … |
要求：每个 Agent 至少 1 个 skill；skill 名称要能对应到代码模块；不要只写「规划/执行/总结」空角色。

【创新点】（≥2 条；必须可编码、可评测；相对传统基线；收益必须写目标百分比）
创新点 1：
- 传统/朴素基线（数值口径）：…
- 你的做法（对应技术栈哪一项 / 哪一流程阶段）：…
- 目标变化：↑xx% 或 ↓xx%（写清相对哪个指标）
- 测量方法：案例集命令 + 真实数据导入后同命令
创新点 2：
- 传统/朴素基线：…
- 你的做法：…
- 目标变化：↑/↓xx%
- 测量方法：…
（可选创新点 3：…）

【指标对比验证 — 证明项目更好】
- 产出 METRICS.md 对照表：指标 | 传统基线 | 改进后 | 变化% | 测法
- STAR-R 必须绑定设计：因为做了…，所以…指标 ↑/↓ xx%
- 未实测写「目标% + 测法」；实测后回填真实数；禁止编造无法复现的大盘数字

【评测数据双通道】
- 案例数据：data/samples/ + 一键评测脚本/按钮
- 真实数据：UI 上传 或 CLI 指定路径；格式/字段 schema；本地处理、不默认外传
- 两种入口共用同一 eval 流水线，都输出百分比对照表

【技术栈硬性清单】
- 核心技术/方法 ≥8；较新栈≥3（飞书手册具体新点+业务理由，禁止空名词与一律硬塞）；自研改进≥2；每项标注流程阶段；标出哪些条目支撑上述创新点
- 必须完成：{task_must_s}
- 范围开关：微调={'开' if do_ft else '关'}；预训练={'开' if do_pt else '关'}；多模态={'开' if do_mm else '关'}；面试题={'开' if do_iv else '关'}；差异化={'开' if do_diff else '关'}；GitHub={'开' if do_gh else '关'}；在线Demo={'开' if do_demo else '关'}
- 禁止：虚构线上 DAU/准确率；无测评集不写结果数字；只堆 Multi-Agent/RAG 空名词不落地；搭建提示词内不要写面试题{'；尽量避免与市面 Demo 同质化' if do_diff else '（差异化关闭：可做客服/问答等常规方向）'}

【强制执行清单 — 按序单任务通关】
{action_s}
0) 接入推理：按 {provider} / {model}（{region_label}）写好客户端、.env.example、失败重试与超时；并落地 SSE 流式

【验收清单】
{result_s}
- [ ] 推理主路径确认为 {provider}/{model}，未擅自换成其他路线默认模型
- [ ] 主路径可演示：LLM 全链路流程跑通 + ↑/↓% 对照表 + 真实数据可导入
从 TODO 第 1 条开始做；先落实流程、LLM_STACK 覆盖表与评测双通道，再写代码。做完一条停下来汇报。
（本块是本项目唯一搭建提示词；禁止追加面试追问或其它提示词）
```

---
{more_projects_full}
---

# ✅ 选题与交付建议（短；不要作品集专章）
- 时间只够做一个时选哪个、为什么
- {'如何避免同质化' if do_diff else '常规方向如何仍讲出工程/算法亮点'}
- 简历一行示例（业务痛点 + 技术栈 + 改进↑/↓%，勿写公司名）
- 自检：每项目 ≤15000 字；每项目恰好 1 份搭建提示词；面试题={'开' if do_iv else '关'}；多模态={'开' if do_mm else '关'}
"""

    return stream_resp(prompt, temperature=0.35, max_tokens=max_tok)


def _assess_project_draft(
    content: str,
    include_finetune: bool = True,
    include_github: bool = True,
    include_demo: bool = True,
    detail_level: str = "detailed",
    project_count: int = 3,
    project_tiers: Optional[List[str]] = None,
    include_pretrain: bool = False,
    include_interview_qa: bool = False,
    include_multimodal: bool = False,
) -> dict:
    """启发式检测项目推荐文稿是否完整（应对输出截断）。"""
    text = (content or "").strip()
    missing = []
    notes = []
    is_brief = (detail_level or "detailed").strip().lower() in ("brief", "short", "简短", "简短版")
    plan = _normalize_project_plan(project_count, project_tiers)
    proj_n = len(plan)
    min_chars = (300 if is_brief else 500) * proj_n // 2 + 200

    if len(text) < min_chars:
        missing.append("正文过短，疑似未开始或严重截断")

    # 按所选数量/质量检测
    found = {}
    for p in plan:
        ok = bool(
            re.search(rf"项目\s*{p['idx']}|{re.escape(p['tag'])}|{re.escape(p['title'])}", text)
        )
        found[p["idx"]] = ok
        if not ok:
            missing.append(f"缺少{p['title']}")

    # 结尾建议
    if not re.search(r"选题与交付建议|交付建议", text):
        missing.append("缺少「选题与交付建议」收尾")

    # AI 提示词：每个项目恰好 1 份
    fence_pairs = text.count("```") // 2
    ai_prompt_hits = len(re.findall(r"交给其他 AI|自动化搭建提示词|精简搭建提示词", text))
    if ai_prompt_hits < proj_n:
        missing.append(f"AI 搭建提示词不足（期望每项目 1 份，共 {proj_n} 份，现约 {ai_prompt_hits}）")
    elif ai_prompt_hits > proj_n + 1:
        notes.append(f"搭建提示词疑似偏多（约 {ai_prompt_hits} 处，期望 {proj_n}）")
    if fence_pairs < proj_n:
        missing.append(f"提示词代码块不足（期望约 {proj_n} 个 ```text）")

    # 面试相关问题 / 作品集附录
    iv_hits = len(re.findall(r"面试相关问题|面试追问|模拟面试|追问预案", text))
    if include_interview_qa:
        if iv_hits < max(1, proj_n):
            missing.append(f"已开启面试相关问题，但未见足够专章（期望约 {proj_n} 处）")
    elif re.search(r"面试追问|模拟面试|追问预案|常见追问", text):
        missing.append("面试相关问题已关闭，出现「面试追问」等内容，应删除")
    if re.search(r"作品集专章|作品集附录|作品集包装|STAR_NARRATIVE", text, flags=re.I):
        notes.append("疑似含作品集附录，建议删掉只保留项目方案与搭建提示词")

    # AI 提示词应含流程 + 创新点 + 百分比验证 + 双通道数据
    flow_hits = len(re.findall(r"端到端流程|【端到端流程】|PIPELINE", text))
    innov_hits = len(re.findall(r"【创新点】|创新点\s*[123一二三]|IMPROVEMENTS", text))
    pct_hits = len(re.findall(r"[↑↓]\s*\d+\s*%|\d+\s*%|[上下]降\s*\d+|提升\s*\d+\s*%|降低\s*\d+\s*%", text))
    dual_hits = len(re.findall(r"真实数据|案例数据|data/samples|上传|导入", text))
    if flow_hits < max(1, proj_n - 1):
        missing.append("AI 搭建提示词缺少「端到端流程」说明")
    if innov_hits < max(1, proj_n - 1):
        missing.append("AI 搭建提示词缺少「创新点」说明")
    if pct_hits < max(2, proj_n):
        missing.append("指标对比缺少↑/↓百分比（STAR-R / 对照表）")
    if dual_hits < max(1, proj_n - 1):
        missing.append("缺少「案例数据 + 真实数据」双通道评测说明")

    # 自研改进 ≥2
    improve_hits = len(re.findall(r"自研改进|创新点\s*[12一二]|改进对照表|IMPROVEMENTS", text))
    if improve_hits < max(1, proj_n):
        missing.append("自研改进/创新点疑似不足（要求每项目 ≥2）")

    # 较新栈：应出现具体新点而非空名词堆砌
    new_stack_hits = len(re.findall(
        r"MCP Apps|Tasks|Streamable HTTP|MRTR|InputRequired|Supervisor|Handoff|Agentic RAG|"
        r"RAG-Fusion|Hybrid|RRF|HyDE|RAPTOR|CRAG|Parent Document|ColBERT|Cohere Rerank|"
        r"ReAct|Plan-and-Execute|Plan and Execute|checkpointer|Human-in-the-loop|interrupt|"
        r"摘要记忆|Harness|投机解码|KV Cache|vLLM|SGLang|"
        r"Local\+Global|社区摘要|层级 GraphRAG|时序图谱|Graphiti|BGE-M3|语义分块",
        text,
        flags=re.I,
    ))
    if new_stack_hits < 1 and re.search(r"较新栈|1\.3b|核心技术栈", text):
        notes.append("较新栈宜写清飞书手册中具体新进展并给业务选型理由（勿只写空名词）")

    # 技术栈表
    stack_hits = len(re.findall(r"核心技术栈|1\.3b|技术栈（≥8）|技术栈\(≥8\)|技术栈（≥7）|技术栈\(≥7\)", text))
    if stack_hits < max(1, min(proj_n, 2)):
        missing.append("核心技术栈（≥8）章节疑似不全")

    # LLM 主流技术覆盖（完整性）
    llm_cov_needles = [
        (r"RAG|Hybrid|RRF|RAG-Fusion|HyDE|RAPTOR|CRAG|Parent Document|ColBERT|Rerank|重排|GraphRAG|图谱", "RAG"),
        ("Multi-Agent|多智能体|多 Agent", "Multi-Agent"),
        ("Function Calling|工具调用|function call", "Function Calling"),
        (r"\bMCP\b|Model Context Protocol", "MCP"),
        ("记忆系统|长期记忆|工作记忆|会话记忆", "记忆系统"),
        ("流式|SSE|streaming", "流式输出"),
        ("RoPE|长上下文|YaRN|位置编码", "RoPE/长上下文"),
        ("后训练|SFT|LoRA|QLoRA|DPO", "后训练"),
    ]
    if include_pretrain:
        llm_cov_needles.append(("预训练|CPT|继续预训练", "预训练"))
    llm_missing = []
    for pat, name in llm_cov_needles:
        if not re.search(pat, text, flags=re.I):
            llm_missing.append(name)
    if len(llm_missing) >= 3:
        missing.append("LLM 主流技术覆盖不足（缺：" + "、".join(llm_missing[:6]) + "）")
    elif llm_missing:
        notes.append("LLM 覆盖表可能不完整（未见：" + "、".join(llm_missing) + "）")
    if not re.search(r"1\.3c|LLM 主流技术覆盖|主流技术覆盖表|LLM_STACK", text, flags=re.I):
        notes.append("建议补上「LLM 主流技术覆盖表 / 1.3c」专章")

    skill_hits = len(re.findall(r"Agent-Skill|AGENT_SKILLS|对应\s*skill|skill\s*映射|task_planning_skill|graph_retrieval_skill|hybrid_retrieval_skill|rerank_skill|mcp_tool_calling_skill|eval_metric_skill", text, flags=re.I))
    if skill_hits < max(1, proj_n):
        missing.append("缺少 Agent-Skill 映射（设计 Agent 时必须搭配对应 skill）")

    if include_finetune and found.get(1) and not re.search(r"微调|后训练|LoRA|QLoRA|SFT", text):
        missing.append("已开启微调，但正文几乎未见后训练/微调方案")
    if include_pretrain and found.get(1) and not re.search(r"预训练|CPT|继续预训练", text):
        missing.append("已开启预训练，但正文几乎未见 CPT/预训练方案")
    if include_multimodal and found.get(1) and not re.search(
        r"多模态|视觉模型|图像理解|音视频|VLM|vision|OCR", text, flags=re.I
    ):
        missing.append("已开启多模态，但正文几乎未见多模态方案")
    if include_github and found.get(1) and not re.search(r"GitHub|仓库|推仓|公开仓", text):
        missing.append("已开启 GitHub，但正文几乎未见仓库/推仓方案")
    if include_demo and found.get(1) and not re.search(r"Demo|公网|在线部署|阿里云|腾讯云|Sealos|魔搭", text):
        missing.append("已开启在线 Demo，但正文几乎未见 Demo/公网方案")

    # STAR：每个所选项目各一套
    sit = len(re.findall(r"###\s*S\s+Situation|Situation（", text, flags=re.I))
    task = len(re.findall(r"###\s*T\s+Task|Task（", text, flags=re.I))
    act = len(re.findall(r"###\s*A\s+Action|Action（", text, flags=re.I))
    res = len(re.findall(r"###\s*R\s+Result|Result（", text, flags=re.I))
    if min(sit, task, act, res) < proj_n:
        missing.append(f"{proj_n} 个项目的完整 STAR（S/T/A/R）未写全，疑似略写或截断")
    if re.search(r"(结构同上|同项目\s*1|见项目\s*1|STAR.*略|略写STAR|见上文.*STAR)", text, flags=re.I):
        missing.append("出现「结构同上/略」等省略写法，STAR 不完整")

    # 截断迹象
    tail = text[-240:]
    open_fences = text.count("```")
    if open_fences % 2 == 1:
        missing.append("存在未闭合的代码块（```），疑似中途截断")
    if re.search(r"(请从|必须写满|【T Task】|【A Action])\s*$", text):
        missing.append("停在提示词模板中间，疑似截断")
    if re.search(r"(##\s*$|#\s*🚀\s*项目\s*\d+\s*[:：]?\s*$)", text):
        missing.append("停在章节标题处，疑似截断")
    if re.search(r"(\.\.\.|…)\s*$", tail) and "选题与交付建议" not in text[-800:]:
        notes.append("末尾像未完成省略，建议续写")

    complete = len(missing) == 0
    return {
        "complete": complete,
        "missing": missing,
        "notes": notes,
        "chars": len(text),
        "project_count": proj_n,
        "project_tiers": [p["tier"] for p in plan],
        "suggestion": (
            "内容看起来完整，可直接复制使用。"
            if complete
            else "内容不完整，建议点击「检测并续写」让 Agent 从断点继续生成。"
        ),
    }


@app.post("/api/project-check")
async def project_check(req: ProjectDraftReq):
    if not (req.content or "").strip():
        return JSONResponse({"error": "暂无项目推荐内容可检测"}, status_code=400)
    do_gh, do_demo = _resolve_deploy_flags(req.include_github, req.include_demo, req.include_deploy)
    result = _assess_project_draft(
        req.content,
        req.include_finetune,
        do_gh,
        do_demo,
        req.detail_level,
        req.project_count,
        req.project_tiers,
        include_pretrain=bool(getattr(req, "include_pretrain", False)),
        include_interview_qa=bool(getattr(req, "include_interview_qa", False)),
        include_multimodal=bool(getattr(req, "include_multimodal", False)),
    )
    return {"ok": True, **result}


@app.post("/api/project-continue")
async def project_continue(req: ProjectDraftReq):
    """从截断处无缝续写项目推荐。"""
    if not (req.content or "").strip():
        return JSONResponse({"error": "没有可续写的草稿"}, status_code=400)
    do_gh, do_demo = _resolve_deploy_flags(req.include_github, req.include_demo, req.include_deploy)
    plan = _normalize_project_plan(req.project_count, req.project_tiers)
    do_pt = bool(getattr(req, "include_pretrain", False))
    do_iv = bool(getattr(req, "include_interview_qa", False))
    do_mm = bool(getattr(req, "include_multimodal", False))
    do_diff = bool(getattr(req, "differentiate", True))
    assessment = _assess_project_draft(
        req.content,
        req.include_finetune,
        do_gh,
        do_demo,
        req.detail_level,
        req.project_count,
        req.project_tiers,
        include_pretrain=do_pt,
        include_interview_qa=do_iv,
        include_multimodal=do_mm,
    )
    missing_s = "、".join(assessment["missing"]) if assessment["missing"] else "可能中途截断"
    draft_tail = req.content[-12000:]
    is_brief = (req.detail_level or "detailed").strip().lower() in ("brief", "short", "简短", "简短版")
    detail_label = "简短版" if is_brief else "详细版"
    tier_s = "、".join(p["tier"] for p in plan)
    _cov_cont = (
        "LLM/RoPE/RAG按业务选型/Multi-Agent/Function Calling/MCP/记忆/流式/部署/后训练"
        + ("/预训练" if do_pt else "（预训练关）")
        + "；禁止一律硬塞 GraphRAG/固定套路，全链路方法按飞书手册+业务选型"
    )
    prompt = f"""你是项目方案架构师。下面是一份「项目推荐」Markdown 草稿，因长度限制被截断，请从断点**无缝续写**。

## 硬性要求
1. **不要重复**已有内容；紧接最后一行继续写。
2. 补全缺失项：{missing_s}
3. 必须最终覆盖：恰好 {len(plan)} 个项目（质量：{tier_s}），以及「选题与交付建议」。
4. 每个项目：{'差异化业务（尽量避免同质化）' if do_diff else '业务方向不限（可含客服/问答等常规方向）'}、无公司名、≥8 技术栈、LLM 主流技术全链路覆盖表（{_cov_cont}）、Agent-Skill 映射表（每个 Agent 搭配对应 skill）、相对传统方案的改进对照（含↑/↓%）、完整 STAR、**恰好 1 份**自动化搭建提示词 ```text。
4b. 每项目整体 ≤15000 字；**禁止**作品集专章/附录、第二份及以上提示词；技术栈必须按工程流程串联，禁止只堆名词；面试题：{'开（须补面试相关问题专章）' if do_iv else '关（不要写面试追问）'}；多模态：{'开（须含非纯文本链路）' if do_mm else '关（勿硬塞多模态）'}。
5. 微调开关：{'开' if req.include_finetune else '关'}；预训练：{'开' if do_pt else '关'}；多模态：{'开' if do_mm else '关'}；面试题：{'开' if do_iv else '关'}；差异化：{'开' if do_diff else '关'}；GitHub：{'开' if do_gh else '关'}；在线Demo：{'开' if do_demo else '关'}（关则不要硬写对应专章；微调关时后训练写「选用已后训练权重」；预训练关时不写 CPT）。
6. 若开启 Demo：仅国内可访问平台（禁 HF Spaces/Vercel/Railway/Streamlit Cloud 等需科学上网）。
7. 篇幅模式：{detail_label}（{'保持精简，勿突然改成长文展开' if is_brief else '按详细版补全缺失小节'}）。
8. 完成时间：每个项目按 {(req.timeline_days if getattr(req, 'timeline_days', None) else 7)} 天规划，不要拉长到默认两周。
9. 项目背景：{('企业级项目' if (req.project_scale or 'school').strip().lower() in ('enterprise', '企业', '企业级', '企业级项目', 'company', 'corp', 'industry') else '学校级项目')}——续写叙事必须与此一致。
9b. 项目形式：{(
        '浏览器插件' if (req.project_form or 'web').strip().lower() in ('extension', 'plugin', 'browser', '浏览器插件', '插件', 'chrome', 'edge')
        else '电脑端 App' if (req.project_form or 'web').strip().lower() in ('desktop', 'pc', '电脑端', '桌面', '桌面端', 'desktop-app', 'electron', 'tauri')
        else '移动端 App' if (req.project_form or 'web').strip().lower() in ('mobile', 'app', '移动端', '移动应用', 'ios', 'android', '手机端')
        else '网站'
    )}——续写必须保持该形态。
10. 只输出续写正文（Markdown），不要开场白。
11. 若缺 STAR / 搭建提示词 / 百分比对照 / LLM 覆盖表 / Agent-Skill 映射表{' / 面试相关问题' if do_iv else ''}{' / 多模态方案' if do_mm else ''}，一并补全；作品集附录不要写；面试题与多模态按开关处理。

## 目标 JD（参考）
{req.jd or '（未提供，按草稿上下文续写）'}

## 已有草稿（末尾最重要，请接上）
{draft_tail}
"""
    return stream_resp(prompt, temperature=0.3, max_tokens=4096 if is_brief else 8192)


@app.get("/api/rag/stats")
async def rag_stats():
    """飞书知识库 RAG 索引状态。"""
    try:
        rag = get_rag()
        rag.ensure_loaded()
        return {"ok": True, **rag.stats()}
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)


@app.post("/api/rag/reload")
async def rag_reload():
    """重新扫描 knowledge/feishu 并重建索引。"""
    try:
        stats = get_rag().reload()
        return {"ok": True, **stats}
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)


@app.post("/api/rag/search")
async def rag_search(req: Request):
    """调试用：对飞书知识库做 BM25 检索。"""
    try:
        body = await req.json()
    except Exception:
        body = {}
    q = (body.get("query") or "").strip()
    top_k = int(body.get("top_k") or 5)
    if not q:
        return JSONResponse({"error": "query required"}, status_code=400)
    hits = get_rag().search(q, top_k=max(1, min(top_k, 12)))
    return {"ok": True, "hits": hits}


@app.post("/api/project-agent")
async def project_agent(req: ProjectDraftReq):
    """项目推荐栏底部 Agent：问答 / 按指令改写 / 主动判断是否需续写。"""
    msg = (req.message or "").strip()
    if not msg:
        return JSONResponse({"error": "请输入要对 Agent 说的话"}, status_code=400)

    do_gh, do_demo = _resolve_deploy_flags(req.include_github, req.include_demo, req.include_deploy)
    do_pt = bool(getattr(req, "include_pretrain", False))
    do_iv = bool(getattr(req, "include_interview_qa", False))
    do_mm = bool(getattr(req, "include_multimodal", False))
    do_diff = bool(getattr(req, "differentiate", True))
    assessment = _assess_project_draft(
        req.content or "",
        req.include_finetune,
        do_gh,
        do_demo,
        req.detail_level,
        req.project_count,
        req.project_tiers,
        include_pretrain=do_pt,
        include_interview_qa=do_iv,
        include_multimodal=do_mm,
    )
    want_continue = any(
        k in msg for k in ("续写", "继续", "补全", "写完", "不完整", "截断", "接着写", "检测")
    )

    if want_continue and not assessment["complete"] and (req.content or "").strip():
        # 走续写逻辑
        return await project_continue(req)

    system = (
        "你是「项目推荐」栏目的协作 Agent。用户在做可落地的面试项目选题。"
        "根据当前草稿回答问题、给修改建议，或按用户要求改某一段。"
        "约束：每项目≤15000字；每项目只有 1 份搭建提示词；不要写作品集附录；"
        f"面试相关问题开关={'开' if do_iv else '关'}；多模态开关={'开' if do_mm else '关'}；"
        "技术栈≥8，须体现主流大模型全链路（LLM/RoPE/RAG按业务选型/Multi-Agent/Function Calling/MCP/记忆/流式/部署/后训练；预训练/多模态按用户开关；禁止一律 GraphRAG），"
        "并按真实工程流程串联；设计 Agent 时必须搭配对应 skill，写清 Agent-Skill 映射表（角色、skill、输入输出、工具/MCP、失败回退、验收指标）。"
        "较新栈≥3、自研改进≥2 比例保持；全链路优先参考飞书手册里较新、效果更好的方法，但**必须按实际项目业务背景选型**。"
        "RAG：原先强制 GraphRAG 已取消；改为业务驱动选合适方案并保证效果，手册中的 Hybrid/RRF、RAPTOR、CRAG 等只是参考示例，**无一强制**。"
        + ("差异化开关=开：尽量避免与市面 Demo 同质化。" if do_diff else "差异化开关=关：可做客服/问答等常规方向，不强制差异化。")
        + "相对传统方案要有可量化改进。"
        "若草稿不完整，主动建议续写并说明缺什么。"
        "回答用中文 Markdown，简洁可执行。"
    )
    _scale = (req.project_scale or "school").strip().lower()
    _scale_label = (
        "企业级项目"
        if _scale in ("enterprise", "企业", "企业级", "企业级项目", "company", "corp", "industry")
        else "学校级项目"
    )
    _form = (req.project_form or "web").strip().lower()
    if _form in ("extension", "plugin", "browser", "浏览器插件", "插件", "chrome", "edge"):
        _form_label = "浏览器插件"
    elif _form in ("desktop", "pc", "电脑端", "桌面", "桌面端", "desktop-app", "electron", "tauri"):
        _form_label = "电脑端 App"
    elif _form in ("mobile", "app", "移动端", "移动应用", "ios", "android", "手机端"):
        _form_label = "移动端 App"
    else:
        _form_label = "网站"
    status = (
        f"完整度：{'完整' if assessment['complete'] else '不完整'}；"
        f"字数：{assessment['chars']}；"
        f"缺失：{('、'.join(assessment['missing']) if assessment['missing'] else '无')}；"
        f"项目背景：{_scale_label}；项目形式：{_form_label}；"
        f"开关：微调={'开' if req.include_finetune else '关'} / "
        f"预训练={'开' if do_pt else '关'} / "
        f"面试题={'开' if do_iv else '关'} / "
        f"多模态={'开' if do_mm else '关'} / "
        f"差异化={'开' if do_diff else '关'} / "
        f"GitHub={'开' if do_gh else '关'} / Demo={'开' if do_demo else '关'}。"
    )
    draft_snip = (req.content or "")[-8000:]

    # RAG：仅增强本 Agent 回复，不改动 /api/project-recommend 生成提示词
    use_rag = bool(getattr(req, "use_rag", True))
    rag_hits: List[dict] = []
    rag_block = ""
    if use_rag:
        try:
            rag = get_rag()
            q = f"{msg}\n{req.jd or ''}".strip()
            rag_hits = rag.search(q, top_k=5)
            rag_block = rag.format_context(rag_hits)
        except Exception:
            rag_hits = []
            rag_block = ""

    if rag_block:
        system += (
            " 回答时可参考下方「知识库检索结果」（来自飞书 RAG/Agent/LangChain 等手册），"
            "优先用其中可落地的技术点增强项目建议；若检索无关则忽略，勿编造手册中没有的内容。"
            "可在回答末尾用一行列出参考了哪些来源文档名。"
        )

    user_prompt = f"""## 当前完整度检测
{status}
建议：{assessment['suggestion']}

## 目标 JD
{req.jd or '（未提供）'}

## 当前草稿（可能很长，已截取尾部+上下文）
{draft_snip or '（尚无生成内容）'}

## 用户消息
{msg}
"""
    if rag_block:
        user_prompt += f"""
## 知识库检索结果（飞书手册 RAG，供参考）
{rag_block}
"""
    messages = [
        SystemMessage(content=system),
        HumanMessage(content=user_prompt),
    ]
    meta = {
        "rag": {
            "enabled": use_rag,
            "hit_count": len(rag_hits),
            "sources": [
                {"source": h["source"], "title": h["title"], "score": h["score"]}
                for h in rag_hits
            ],
        }
    }
    return stream_resp(messages, temperature=0.4, max_tokens=4096, meta=meta)


def _md_plain_lines(content: str) -> List[str]:
    """把 Markdown 粗略拆成适合写入 Word/PDF 的行。"""
    lines = []
    in_code = False
    for raw in (content or "").replace("\r\n", "\n").split("\n"):
        line = raw.rstrip()
        if line.strip().startswith("```"):
            in_code = not in_code
            lines.append(line)
            continue
        lines.append(line)
    return lines


def _export_docx_bytes(title: str, content: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn

    def set_cn_font(run, name: str = "Microsoft YaHei", size: Optional[int] = None):
        """同时设置西文字体和东亚字体，避免 Word 中文显示乱码/回退异常。"""
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        if size:
            run.font.size = Pt(size)

    doc = Document()
    h = doc.add_heading(title or "项目推荐方案", level=0)
    for run in h.runs:
        set_cn_font(run, size=18)

    in_code = False
    for line in _md_plain_lines(content):
        s = line.strip()
        if s.startswith("```"):
            in_code = not in_code
            continue
        if not s:
            doc.add_paragraph("")
            continue
        if not in_code and s.startswith("#"):
            level = min(len(s) - len(s.lstrip("#")), 3)
            text = s.lstrip("#").strip()
            if text:
                heading = doc.add_heading(text, level=level)
                for run in heading.runs:
                    set_cn_font(run)
            continue
        bullet = not in_code and re.match(r"^[-*]\s+", s)
        text = re.sub(r"^[-*]\s+", "", s) if bullet else line
        p = doc.add_paragraph(style="List Bullet" if bullet else None)
        # 支持 Markdown 粗体，同时确保每个 run 使用中文字体。
        parts = re.split(r"(\*\*.+?\*\*)", text)
        for part in parts:
            if not part:
                continue
            bold = part.startswith("**") and part.endswith("**")
            run = p.add_run(part[2:-2] if bold else part)
            run.bold = bold
            set_cn_font(run)
        if in_code:
            for run in p.runs:
                set_cn_font(run, "Microsoft YaHei", 9)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for section in doc.sections:
        section.top_margin = Pt(42)
        section.bottom_margin = Pt(42)
        section.left_margin = Pt(48)
        section.right_margin = Pt(48)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _find_cn_font() -> Optional[str]:
    candidates = [
        r"C:\Windows\Fonts\msyh.ttf",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\Deng.ttf",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    ]
    for p in candidates:
        if Path(p).exists():
            return p
    return None


def _export_pdf_bytes(title: str, content: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted

    font_path = _find_cn_font()
    font_name = "Helvetica"
    if font_path:
        try:
            # Windows 常见中文字体可能是 TTC；显式选择首个子字体。
            pdfmetrics.registerFont(TTFont("CNFont", font_path, subfontIndex=0))
            font_name = "CNFont"
        except Exception:
            # Helvetica 不含中文字形，不能静默回退，否则导出的 PDF 会乱码。
            raise RuntimeError("未能加载中文字体，请确认系统已安装微软雅黑、黑体或等线字体")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.8 * cm, rightMargin=1.8 * cm,
        topMargin=1.8 * cm, bottomMargin=1.8 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CNTitle", parent=styles["Title"], fontName=font_name, fontSize=16, leading=22, spaceAfter=12
    )
    h1 = ParagraphStyle("CNH1", parent=styles["Heading1"], fontName=font_name, fontSize=14, leading=20, spaceBefore=10, spaceAfter=6)
    h2 = ParagraphStyle("CNH2", parent=styles["Heading2"], fontName=font_name, fontSize=12, leading=18, spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("CNBody", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=15, spaceAfter=4)
    code = ParagraphStyle("CNCode", parent=styles["Code"], fontName=font_name, fontSize=8, leading=11, spaceAfter=4)

    story = [Paragraph(_xml_escape(title or "项目推荐方案"), title_style), Spacer(1, 8)]
    in_code = False
    code_buf = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            story.append(Preformatted("\n".join(code_buf), code))
            code_buf = []

    for line in _md_plain_lines(content):
        s = line.strip()
        if s.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        if not s:
            story.append(Spacer(1, 4))
            continue
        if s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            text = _xml_escape(s.lstrip("#").strip())
            story.append(Paragraph(text, h1 if level <= 1 else h2))
            continue
        # 简单处理粗体 **x**
        html = _xml_escape(s)
        html = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", html)
        story.append(Paragraph(html, body))

    if in_code:
        flush_code()

    doc.build(story)
    return buf.getvalue()


def _xml_escape(s: str) -> str:
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


@app.post("/api/export-doc")
async def export_doc(req: ExportDocReq):
    """将 Markdown 正文导出为 Word 或 PDF。"""
    content = (req.content or "").strip()
    if not content:
        return JSONResponse({"error": "没有可导出的内容"}, status_code=400)
    fmt = (req.format or "docx").strip().lower()
    title = (req.title or "项目推荐方案").strip() or "项目推荐方案"
    try:
        if fmt in ("docx", "word", "doc"):
            data = _export_docx_bytes(title, content)
            filename = f"{title}.docx"
            media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif fmt == "pdf":
            data = _export_pdf_bytes(title, content)
            filename = f"{title}.pdf"
            media = "application/pdf"
        else:
            return JSONResponse({"error": "format 仅支持 docx 或 pdf"}, status_code=400)
    except Exception as e:
        return JSONResponse({"error": f"导出失败：{e}"}, status_code=500)

    # 文件名兼容中文
    from urllib.parse import quote
    encoded = quote(filename)
    return StreamingResponse(
        io.BytesIO(data),
        media_type=media,
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
            "Content-Length": str(len(data)),
        },
    )


@app.post("/api/chat")
async def chat(req: ChatReq):
    messages = [SystemMessage(content=(
        "你是专业的求职顾问，帮助用户提升求职成功率。"
        "回答简洁专业，使用中文，给出可执行的具体建议。"
    ))]
    for item in req.history[-10:]:
        if item.role == "user":
            messages.append(HumanMessage(content=item.content))
        else:
            messages.append(AIMessage(content=item.content))
    messages.append(HumanMessage(content=req.message))
    return stream_resp(messages)

@app.post("/api/parse-resume-image")
async def parse_resume_image(file: UploadFile = File(...)):
    """从简历图片中提取文字（需要视觉模型，如 gpt-4o / deepseek-vl）"""
    allowed = {"image/jpeg", "image/png", "image/webp", "image/gif"}
    if file.content_type not in allowed:
        return JSONResponse({"error": "仅支持 JPG / PNG / WEBP / GIF 图片"}, status_code=400)

    data = await file.read()
    if len(data) > 10 * 1024 * 1024:
        return JSONResponse({"error": "图片不能超过 10 MB"}, status_code=400)

    b64 = base64.b64encode(data).decode()
    mime = file.content_type

    message = HumanMessage(content=[
        {
            "type": "image_url",
            "image_url": {"url": f"data:{mime};base64,{b64}"},
        },
        {
            "type": "text",
            "text": (
                "请识别并提取这张简历图片中的全部文字内容。"
                "保持原有层级结构（姓名、联系方式、教育背景、工作经历、技能等），"
                "用纯文本输出，不要添加任何额外说明或 Markdown 标记。"
            ),
        },
    ])

    async def generate():
        try:
            async for chunk in get_llm().astream([message]):
                if chunk.content:
                    d = json.dumps({"content": chunk.content}, ensure_ascii=False)
                    yield f"data: {d}\n\n"
        except Exception as e:
            err = json.dumps({"error": str(e)}, ensure_ascii=False)
            yield f"data: {err}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )

@app.post("/api/parse-resume-file")
async def parse_resume_file(file: UploadFile = File(...)):
    """从 Word (.docx) 或 PDF 文件中直接提取文字，保留结构排版"""
    fname = (file.filename or "").lower()
    data = await file.read()
    if len(data) > 20 * 1024 * 1024:
        return JSONResponse({"error": "文件不能超过 20 MB"}, status_code=400)

    text = ""
    try:
        if fname.endswith(".docx"):
            import docx
            from docx.text.paragraph import Paragraph
            from docx.table import Table as DocxTable

            doc = docx.Document(io.BytesIO(data))
            W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            lines = []
            seen = set()

            def add_line(t, is_heading=False):
                t = t.strip()
                if not t or t in seen:
                    return
                seen.add(t)
                if is_heading:
                    lines.append("")
                    lines.append(f"【{t}】")
                else:
                    lines.append(t)

            def proc_tbl(el):
                tbl = DocxTable(el, doc)
                for row in tbl.rows:
                    cells = []
                    prev = None
                    for c in row.cells:
                        ct = c.text.strip()
                        if ct and ct != prev:
                            cells.append(ct)
                            prev = ct
                    if cells:
                        add_line("  ".join(cells))

            def proc_para(el):
                para = Paragraph(el, doc)
                t = para.text.strip()
                if not t:
                    return
                style = para.style.name if para.style else ""
                is_hd = any(k in style for k in ("Heading", "heading", "标题", "Title"))
                add_line(t, is_heading=is_hd)

            def walk(el):
                tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
                if tag == "p":
                    proc_para(el)
                    # 深入检查段落内的文本框（mc:AlternateContent 结构）
                    for desc in el.iter():
                        dtag = desc.tag.split("}")[-1] if "}" in desc.tag else desc.tag
                        if dtag == "txbxContent":
                            for child in desc:
                                walk(child)
                elif tag == "tbl":
                    proc_tbl(el)
                elif tag == "txbxContent":
                    for child in el:
                        walk(child)
                else:
                    for child in el:
                        walk(child)

            for child in doc.element.body:
                walk(child)

            text = "\n".join(lines).strip()

        elif fname.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text(layout=True)
                    if not t:
                        t = page.extract_text()
                    if t:
                        pages.append(t.strip())
            text = "\n\n".join(pages)
        else:
            return JSONResponse({"error": "仅支持 .docx 和 .pdf 文件"}, status_code=400)
    except Exception as e:
        return JSONResponse({"error": f"文件解析失败：{e}"}, status_code=500)

    if not text.strip():
        return JSONResponse({"error": "未能提取到文字，请确认文件内容不是纯图片扫描版"}, status_code=422)

    return JSONResponse({"text": text.strip()})


@app.post("/api/run-code")
async def run_code(request: Request):
    """在本机 Python 环境中执行用户代码。公网默认关闭（ENABLE_RUN_CODE=0）。"""
    if not ENABLE_RUN_CODE:
        raise HTTPException(
            status_code=403,
            detail="代码执行已禁用（公网安全策略）。本地可设 ENABLE_RUN_CODE=1",
        )
    body = await request.json()
    user_code = body.get("code", "").strip()
    setup_code = body.get("setup_code", "").strip()
    mode = body.get("mode", "leetcode")  # "leetcode" | "scratch"

    if not user_code:
        return JSONResponse({"error": "代码不能为空"}, status_code=400)

    if mode == "scratch":
        # 手打模式：直接运行用户代码，捕获 stdout/stderr
        full_code = (
            "import sys\n"
            "from typing import List, Optional, Dict, Tuple, Set\n\n"
            + user_code
        )
    else:
        # LeetCode 模式：包裹 setup_code + Solution + _run_tests
        full_code = (
            "import sys, json\n"
            "from typing import List, Optional, Dict, Tuple, Set\n\n"
            + setup_code
            + "\n\n"
            + user_code
            + "\n\ntry:\n"
            "    _r = _run_tests(Solution())\n"
            "    print(json.dumps(_r, ensure_ascii=False))\n"
            "except Exception as _e:\n"
            "    print(json.dumps([{\"case\":1,\"passed\":False,\"input\":\"\",\"error\":str(_e)}]))\n"
        )

    tmpfile = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".py", delete=False, encoding="utf-8"
        ) as f:
            f.write(full_code)
            tmpfile = f.name

        result = subprocess.run(
            [sys.executable, tmpfile],
            capture_output=True, text=True, timeout=10, encoding="utf-8"
        )

        stdout = result.stdout.strip()
        stderr = result.stderr.strip()

        if mode == "scratch":
            # 手打模式：returncode == 0 且无 stderr 视为成功
            success = (result.returncode == 0)
            return JSONResponse({
                "mode": "scratch",
                "success": success,
                "stdout": stdout,
                "stderr": stderr[-600:] if stderr else "",
            })

        # LeetCode 模式
        if not stdout and stderr:
            return JSONResponse([{"case": 1, "passed": False, "input": "",
                                   "error": stderr[-300:]}])
        if not stdout:
            return JSONResponse([{"case": 1, "passed": False, "input": "",
                                   "error": "程序无输出"}])

        return JSONResponse(json.loads(stdout))

    except subprocess.TimeoutExpired:
        if mode == "scratch":
            return JSONResponse({"mode": "scratch", "success": False, "stdout": "",
                                  "stderr": "执行超时（超过 10 秒），请检查是否有死循环"})
        return JSONResponse([{"case": 1, "passed": False, "input": "",
                               "error": "执行超时（超过 10 秒），请检查是否有死循环"}])
    except Exception as e:
        if mode == "scratch":
            return JSONResponse({"mode": "scratch", "success": False, "stdout": "",
                                  "stderr": f"系统错误：{e}"})
        return JSONResponse([{"case": 1, "passed": False, "input": "",
                               "error": f"系统错误：{e}"}])
    finally:
        if tmpfile:
            try:
                os.unlink(tmpfile)
            except:
                pass


# ══════════════════════════════════════════════════════════════
# 💰 支付系统
# ══════════════════════════════════════════════════════════════

PAY_DIR = BASE_DIR / "data"
PAY_DIR.mkdir(exist_ok=True)
PAY_ORDERS_FILE = PAY_DIR / "pay_orders.json"
PAY_USERS_FILE  = PAY_DIR / "pay_users.json"
PAY_CONFIG_FILE = PAY_DIR / "pay_config.json"

PLANS = {
    "day":   {"name": "日卡",  "days": 1,   "amount": 10},
    "month": {"name": "月卡",  "days": 30,  "amount": 200},
    "year":  {"name": "年卡",  "days": 365, "amount": 1000},
}

def _load(path, default):
    if not path.exists():
        return default() if callable(default) else default
    return json.loads(path.read_text(encoding="utf-8"))

def _save(path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def load_pay_config():
    defaults = {"admin_password": "admin123", "alipay_qr": "", "wechat_qr": ""}
    cfg = _load(PAY_CONFIG_FILE, dict)
    return {**defaults, **cfg}

# ── 请求模型 ───────────────────────────────────────────────
class CreateOrderReq(BaseModel):
    plan: str
    email: str
    name: str = ""
    pay_method: str = "alipay"

class SubmitPayReq(BaseModel):
    order_id: str

class CheckLicenseReq(BaseModel):
    email: str

class AdminActReq(BaseModel):
    order_id: str
    password: str

class PayConfigSaveReq(BaseModel):
    admin_password: str
    new_password: str = ""
    alipay_qr: str = ""
    wechat_qr: str = ""

# ── 用户端接口 ─────────────────────────────────────────────
@app.post("/api/pay/create")
async def pay_create(req: CreateOrderReq):
    if req.plan not in PLANS:
        raise HTTPException(status_code=400, detail="无效套餐")
    email = req.email.strip().lower()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="请填写有效邮箱")
    plan = PLANS[req.plan]
    order_id = str(uuid.uuid4())[:8].upper()
    order = {
        "id": order_id,
        "plan": req.plan,
        "plan_name": plan["name"],
        "email": email,
        "name": req.name.strip(),
        "amount": plan["amount"],
        "pay_method": req.pay_method,
        "status": "pending_payment",
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "submitted_at": None,
        "activated_at": None,
        "expires_at": None,
    }
    orders = _load(PAY_ORDERS_FILE, list)
    orders.append(order)
    _save(PAY_ORDERS_FILE, orders)
    cfg = load_pay_config()
    return {
        "order_id": order_id,
        "amount": plan["amount"],
        "plan_name": plan["name"],
        "alipay_qr": cfg.get("alipay_qr", ""),
        "wechat_qr": cfg.get("wechat_qr", ""),
    }

@app.post("/api/pay/submit")
async def pay_submit(req: SubmitPayReq):
    orders = _load(PAY_ORDERS_FILE, list)
    for o in orders:
        if o["id"] == req.order_id:
            o["status"] = "pending_review"
            o["submitted_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            _save(PAY_ORDERS_FILE, orders)
            return {"ok": True}
    raise HTTPException(status_code=404, detail="订单不存在")

@app.post("/api/pay/check")
async def pay_check(req: CheckLicenseReq):
    users = _load(PAY_USERS_FILE, dict)
    email = req.email.strip().lower()
    if email not in users:
        return {"active": False}
    u = users[email]
    expires_at = u.get("expires_at")
    if expires_at:
        exp_dt = datetime.strptime(expires_at, "%Y-%m-%d %H:%M:%S")
        if datetime.now() > exp_dt:
            return {"active": False, "expired": True, "expires_at": expires_at}
    return {
        "active": True,
        "plan": u.get("plan"),
        "expires_at": expires_at,
        "name": u.get("name", ""),
    }

# ── 管理员接口 ─────────────────────────────────────────────
def _check_admin_password(password: str) -> None:
    cfg = load_pay_config()
    if password != cfg["admin_password"]:
        raise HTTPException(status_code=403, detail="密码错误")

@app.get("/admin", response_class=HTMLResponse)
async def admin_page(request: Request):
    if PUBLIC_MODE:
        _require_admin_token(request)
    p = STATIC_DIR / "admin.html"
    if p.exists():
        return HTMLResponse(content=p.read_text(encoding="utf-8"))
    return HTMLResponse("<h2>请先创建 static/admin.html</h2>")

@app.get("/api/admin/orders")
async def admin_orders(request: Request, password: str = ""):
    _require_admin_token(request, password)
    if not PUBLIC_MODE:
        _check_admin_password(password)
    return _load(PAY_ORDERS_FILE, list)

@app.post("/api/admin/activate")
async def admin_activate(request: Request, req: AdminActReq):
    _require_admin_token(request, req.password)
    if not PUBLIC_MODE:
        _check_admin_password(req.password)
    orders = _load(PAY_ORDERS_FILE, list)
    users  = _load(PAY_USERS_FILE, dict)
    for o in orders:
        if o["id"] == req.order_id:
            plan = PLANS.get(o["plan"])
            if not plan:
                raise HTTPException(status_code=400, detail="套餐数据异常")
            now = datetime.now()
            # 如果已有有效期则续期
            email = o["email"]
            existing = users.get(email, {}).get("expires_at")
            base_dt = now
            if existing:
                try:
                    exp_dt = datetime.strptime(existing, "%Y-%m-%d %H:%M:%S")
                    if exp_dt > now:
                        base_dt = exp_dt
                except Exception:
                    pass
            expires = base_dt + timedelta(days=plan["days"])
            o["status"]       = "active"
            o["activated_at"] = now.strftime("%Y-%m-%d %H:%M:%S")
            o["expires_at"]   = expires.strftime("%Y-%m-%d %H:%M:%S")
            users[email] = {
                "name":         o.get("name", ""),
                "plan":         o["plan"],
                "expires_at":   expires.strftime("%Y-%m-%d %H:%M:%S"),
                "activated_at": now.strftime("%Y-%m-%d %H:%M:%S"),
            }
            _save(PAY_ORDERS_FILE, orders)
            _save(PAY_USERS_FILE,  users)
            return {"ok": True, "expires_at": expires.strftime("%Y-%m-%d %H:%M:%S")}
    raise HTTPException(status_code=404, detail="订单不存在")

@app.post("/api/admin/reject")
async def admin_reject(request: Request, req: AdminActReq):
    _require_admin_token(request, req.password)
    if not PUBLIC_MODE:
        _check_admin_password(req.password)
    orders = _load(PAY_ORDERS_FILE, list)
    for o in orders:
        if o["id"] == req.order_id:
            o["status"] = "rejected"
            _save(PAY_ORDERS_FILE, orders)
            return {"ok": True}
    raise HTTPException(status_code=404, detail="订单不存在")

@app.get("/api/admin/pay-config")
async def admin_get_pay_config(request: Request, password: str = ""):
    _require_admin_token(request, password)
    if not PUBLIC_MODE:
        _check_admin_password(password)
    cfg = load_pay_config()
    # 公网不回传管理员密码明文
    if PUBLIC_MODE:
        return {
            "alipay_qr": cfg.get("alipay_qr", ""),
            "wechat_qr": cfg.get("wechat_qr", ""),
            "has_password": bool(cfg.get("admin_password")),
        }
    return cfg

@app.post("/api/admin/pay-config")
async def admin_save_pay_config(request: Request, req: PayConfigSaveReq):
    _require_admin_token(request, req.admin_password)
    if not PUBLIC_MODE:
        _check_admin_password(req.admin_password)
    cfg = load_pay_config()
    if req.new_password.strip():
        cfg["admin_password"] = req.new_password.strip()
    if req.alipay_qr.strip():
        cfg["alipay_qr"] = req.alipay_qr.strip()
    if req.wechat_qr.strip():
        cfg["wechat_qr"] = req.wechat_qr.strip()
    _save(PAY_CONFIG_FILE, cfg)
    return {"ok": True}

